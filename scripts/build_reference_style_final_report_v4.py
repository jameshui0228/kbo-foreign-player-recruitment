#!/usr/bin/env python3
"""Build the reference-style final DOCX report for the SSG recruitment project.

The target format mirrors the provided season-1 SDA PDF: cover, abstract,
key takeaways, contents, table/figure lists, numbered sections, RQ-driven
results, discussion, limitations, conclusion, appendix, and references.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "outputs" / "tables"
REPORTS = ROOT / "outputs" / "reports"
ASSETS = REPORTS / "ensemble_final_report_v1_assets"
OUT = ROOT / "output" / "doc"
DOCX = OUT / "ssg_foreign_player_ensemble_final_report_v4_reference_style.docx"

FONT = "Apple SD Gothic Neo"
INK = "1F2430"
MUTED = "666E80"
LIGHT = "F4F6FA"
BLUE_LIGHT = "EAF1FE"
GOLD_LIGHT = "FFF3C4"
RED_LIGHT = "FDECEC"
GREEN_LIGHT = "EAF7ED"
GRID = "D8DDE8"


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, bold: bool = False, color: str = INK, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)


def add_run(paragraph, text: str, *, size: float | None = None, bold: bool = False, color: str = INK, italic: bool = False):
    r = paragraph.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return r


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border(cell, color: str = GRID, size: str = "5") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def cell_text(cell, text: str, *, size: float = 7.4, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, color: str = INK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(
    doc: Document,
    df: pd.DataFrame,
    columns: list[str],
    headers: list[str],
    caption: str | None = None,
    *,
    size: float = 7.2,
    header_fill: str = LIGHT,
    max_rows: int | None = None,
) -> None:
    if max_rows is not None:
        df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, header_fill)
        border(cell)
        cell_text(cell, header, size=size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for cell, col in zip(cells, columns):
            val = row[col]
            if pd.isna(val):
                txt = "-"
            elif isinstance(val, float):
                txt = f"{val:.3f}" if abs(val) < 10 else f"{val:,.1f}"
            else:
                txt = str(val)
            border(cell)
            cell_text(cell, txt, size=size)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cap, caption, size=8.2, color=MUTED)


def add_p(doc: Document, text: str = "", *, size: float = 9.8, first_line: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.42)
    add_run(p, text, size=size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    add_run(p, text, size=15.5 if level == 1 else 12.2, bold=True)


def add_bullets(doc: Document, rows: list[str]) -> None:
    for row in rows:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_run(p, row, size=9.4)


def add_callout(doc: Document, title: str, body: str, *, fill: str = BLUE_LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    border(cell)
    p = cell.paragraphs[0]
    add_run(p, title, bold=True, size=9.4)
    add_run(p, " " + body, size=9.2)


def add_picture(doc: Document, path: Path, caption: str, *, width: float = 6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        p.add_run().add_picture(str(path), width=Inches(width))
    else:
        add_run(p, f"[그림 파일 없음: {path.name}]", size=8.5, color=MUTED)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, size=8.2, color=MUTED)


def start_page(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)


def set_document_style(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)
    styles = doc.styles
    for name, size, bold in [("Normal", 9.8, False), ("Heading 1", 15.5, True), ("Heading 2", 12.2, True), ("Title", 25, True)]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = rgb(INK)
        if name.startswith("Heading"):
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(5)


def add_footer(doc: Document) -> None:
    for i, sec in enumerate(doc.sections):
        sec.footer.is_linked_to_previous = False
        p = sec.footer.paragraphs[0]
        for run in list(p.runs):
            p._p.remove(run._r)
        if i == 0:
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "SSG 외국인 선수 영입 앙상블 데이터 마이닝 보고서", size=8, color=MUTED)


def fmt_pct(v: float | int | None) -> str:
    if v is None or pd.isna(v):
        return "-"
    return f"{float(v) * 100:.1f}%"


def money(v) -> str:
    if pd.isna(v):
        return "-"
    return f"${int(v):,}"


def prep_tables() -> dict[str, pd.DataFrame]:
    scores = read_csv("ensemble_candidate_scores_v1.csv")
    hitter_dm = read_csv("data_mining_hitter_candidates_v1.csv")
    pitcher_dm = read_csv("data_mining_pitcher_candidates_v1.csv")
    salary_gate = read_csv("final_candidate_salary_contract_gate_v1.csv")
    coverage = read_csv("data_coverage_by_league_v1.csv")
    weights = read_csv("ensemble_model_signal_weights_v1.csv")
    model_audit = read_csv("data_mining_model_audit_v1.csv")
    trace = read_csv("ssg_layer1_evidence_to_message_trace_v4.csv")
    feature_blueprint = read_csv("ssg_layer1_candidate_feature_blueprint_v4.csv")
    layer5 = read_csv("layer5_failure_risk_v0_3_slot_summary_v0_1.csv")
    layer6 = read_csv("layer6_fit_ranking_v0_3_stage_summary_v0_1.csv")

    hidden_rules = pd.DataFrame(
        [
            ["R1", "RHP + OF low conversion", 20, "0.100", "-5.10", "외야/DH가 우투 선발전에서 이닝을 다시 열지 못하는 구간"],
            ["R2", "RHP + OF/DH run-kill", 9, "0.000", "-5.11", "병살/도루자형 단절이 붙는 순간 복구 루트가 사라지는 구간"],
            ["R3", "Extra-out + OF void", 8, "0.000", "-4.50", "비자책/추가 아웃 손실 뒤 추가실점이 커지는 구간"],
            ["R4", "Top opponent + short start + OF silence", 6, "0.000", "-5.83", "강팀전에서 선발 길이와 외야 전환이 동시에 사라진 구간"],
        ],
        columns=["rule", "interaction", "games", "win_pct", "run_diff", "interpretation"],
    )

    team_lens = pd.DataFrame(
        [
            ["sewon", "시장/등급형 후보 보드", "MLB노출·등급·시장 접근성", "Nolan Jones, Jack Suwinski / Carson Spiers, Brian Van Belle", "모델이 놓칠 수 있는 현실적 후보군을 넓힘"],
            ["jimini", "SSG fit/역할 번역 렌즈", "SSG 외야·투수 슬롯의 역할 적합성", "Nolan Jones, Dominic Fletcher / Josh Fleming, Ian Hamilton", "팀 약점과 선수 기능의 연결을 강화"],
            ["kyuho", "Statcast upside/비효율 렌즈", "장타·구위·시장 저평가 가능성", "Jack Suwinski, Michael Toglia / Josh Fleming, Randy Dobnak", "고위험 고보상 후보를 보드에 남김"],
            ["codex", "supervised data-mining lens", "KBO 외인 성공/실패, MiLB/Savant, salary gate", "Luis Matos, Nolan Jones / Josh Fleming, Bryse Wilson", "과거 KBO 패턴과 현재 시장 현실성을 결합"],
        ],
        columns=["analyst", "model_lens", "main_signal", "main_candidates", "ensemble_role"],
    )

    model_blocks = pd.DataFrame(
        [
            ["M1", "SSG hidden-weakness interaction mining", "KBO/STATIZ 상황별 팀 성과", "상대 우투 선발, OF/DH 전환, run-kill, extra-out", "SSG가 필요로 하는 선수 기능을 정의"],
            ["M2", "KBO foreign-player success/failure classifier", "과거 KBO 외인 + pre-KBO Savant/MiLB", "Ridge logistic, sparse L1 logistic", "후보가 KBO 성공 패턴에 가까운지 계산"],
            ["M3", "Candidate market construction", "MLB/MiLB/NPB/CPBL roster, transaction, stats", "40인/비40인, DFA, MiLB role, NPB official stats", "실제 후보 시장을 과대상상하지 않게 제한"],
            ["M4", "KBO translation/adaptation filter", "Savant plate discipline, MiLB command/damage", "RHP 대응, chase/whiff, BB/HR 억제, starter floor", "MLB 성적이 KBO에서 번역될 가능성을 점검"],
            ["M5", "Failure risk hard gate", "medical, injury status, sample, current role", "medical red, tiny sample, lower-level risk", "점수는 좋아도 실패 가능성이 큰 후보를 hold"],
            ["M6", "Salary/contract feasibility gate", "public salary, assignment, minor deal, buyout proxy", "salary signal, 40-man, minor deal, outright/DFA", "최종 후보를 실제 접촉 가능한 보드로 변환"],
        ],
        columns=["model", "name", "data", "features", "role"],
    )

    hitter_names = ["Nolan Jones", "Luis Matos", "Jack Suwinski"]
    hitter_top = scores[scores["candidate"].isin(hitter_names)].copy()
    hitter_top["success_prob"] = hitter_top["dm_success_prob"].map(fmt_pct)
    hitter_top["failure_prob"] = hitter_top["dm_failure_prob"].map(fmt_pct)
    hitter_top["ensemble_score_fmt"] = hitter_top["ensemble_score"].map(lambda x: f"{x:.3f}")
    hitter_top = hitter_top.sort_values("rank")

    pitcher_raw_names = ["Josh Fleming", "Carson Spiers", "Brian Van Belle"]
    pitcher_raw = scores[scores["candidate"].isin(pitcher_raw_names)].copy()
    pitcher_raw["ensemble_score_fmt"] = pitcher_raw["ensemble_score"].map(lambda x: f"{x:.3f}")
    pitcher_raw = pitcher_raw.sort_values("rank")

    pitcher_gate_names = ["Josh Fleming", "Bryse Wilson", "Austin Gomber"]
    pitcher_gate = pitcher_dm[pitcher_dm["player_name"].isin(pitcher_gate_names)].copy()
    pitcher_gate["success_prob"] = pitcher_gate["dm_success_prob"].map(fmt_pct)
    pitcher_gate["failure_prob"] = pitcher_gate["dm_failure_prob"].map(fmt_pct)
    pitcher_gate["margin_fmt"] = pitcher_gate["dm_margin"].map(lambda x: f"{x:.3f}" if pd.notna(x) else "-")
    pitcher_gate = pitcher_gate.sort_values(["data_mining_gate_pass", "dm_margin"], ascending=[False, False])

    hitter_dm_view = hitter_dm[hitter_dm["player_name"].isin(["Luis Matos", "Nolan Jones", "Dylan Carlson", "Jack Suwinski"])].copy()
    hitter_dm_view["success_prob"] = hitter_dm_view["dm_success_prob"].map(fmt_pct)
    hitter_dm_view["failure_prob"] = hitter_dm_view["dm_failure_prob"].map(fmt_pct)
    hitter_dm_view["margin_fmt"] = hitter_dm_view["dm_margin"].map(lambda x: f"{x:.3f}" if pd.notna(x) else "-")
    hitter_dm_view = hitter_dm_view.sort_values("data_mining_rank", na_position="last")

    salary_gate["salary_fmt"] = salary_gate["structured_salary_signal_usd"].map(money)
    final_decision = pd.DataFrame(
        [
            ["외국인 타자", "raw ensemble", "Nolan Jones", "모델 종합 1위", "성공확률 90.2%, consensus 1.00"],
            ["외국인 타자", "salary-gated", "Luis Matos", "최종 접촉 1순위", "DM rank 1, 성공확률 92.4%, 낮은 salary signal"],
            ["외국인 투수", "raw ensemble", "Josh Fleming", "raw/실행 모두 1위", "세 모델 lens에서 반복 등장"],
            ["외국인 투수", "salary/medical-gated", "Josh Fleming", "최종 접촉 1순위", "minor deal 접근성, Spiers/Van Belle medical red"],
        ],
        columns=["slot", "board", "final_one", "decision", "reason"],
    )

    references = pd.DataFrame(
        [
            ["KBO 규정", "외국인/아시아쿼터/대체 외인 비용 gate", "https://www.koreabaseball.com/Kbo/League/GameManage2026.aspx"],
            ["KBO 공지", "대체 외국인 월 비용 제한", "https://www.koreabaseball.com/MediaNews/Notice/View.aspx?bdSe=9984"],
            ["Spotrac", "Nolan Jones 2026 cash signal", "https://www.spotrac.com/mlb/cleveland-guardians/cash"],
            ["MLB.com", "Luis Matos transaction/status", "https://www.mlb.com/player/luis-matos-682641"],
            ["MLB.com", "Jack Suwinski one-year deal", "https://www.mlb.com/news/jack-suwinski-agrees-to-one-year-deal-with-pirates"],
            ["MLB.com", "Josh Fleming transaction/status", "https://www.mlb.com/player/josh-fleming-676596"],
            ["MLBTR", "Carson Spiers minor deal and injury", "https://www.mlbtraderumors.com/2025/11/reds-re-sign-carson-spiers-to-minor-league-deal.html"],
            ["MLB.com", "Brian Van Belle full-season IL", "https://www.mlb.com/player/brian-van-belle-687003"],
        ],
        columns=["source", "use", "url"],
    )

    return {
        "scores": scores,
        "coverage": coverage,
        "weights": weights,
        "model_audit": model_audit,
        "trace": trace,
        "feature_blueprint": feature_blueprint,
        "layer5": layer5,
        "layer6": layer6,
        "hidden_rules": hidden_rules,
        "team_lens": team_lens,
        "model_blocks": model_blocks,
        "hitter_top": hitter_top,
        "pitcher_raw": pitcher_raw,
        "pitcher_gate": pitcher_gate,
        "hitter_dm_view": hitter_dm_view,
        "salary_gate": salary_gate,
        "final_decision": final_decision,
        "references": references,
    }


def cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "SSG 외국인 선수 영입 전략", size=25, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "앙상블 데이터 마이닝 기반 후보 선정 보고서", size=17, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Data: KBO/STATIZ | MLB/MiLB | Savant | NPB/CPBL | Roster/Transaction | Salary Gate", size=9.6)
    for _ in range(13):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "고려대학교 SDA 2기", size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "June 23, 2026", size=10.5)
    start_page(doc)


def abstract(doc: Document) -> None:
    add_heading(doc, "초록 (Abstract)", 1)
    add_heading(doc, "국문", 2)
    add_p(
        doc,
        "본 보고서는 SSG 랜더스의 대체 외국인 타자 및 외국인 선발투수 영입을 단순 후보 추천이 아니라, 팀 약점 마이닝과 시장 현실성 gate를 결합한 데이터 마이닝 문제로 재정의한다. 분석의 출발점은 SSG의 숨은 약점이다. 2026 SSG는 외야 장타 자체보다 상대 우투 선발전에서 OF/DH 전환이 막히고, 병살·도루자형 run-kill이나 비자책/추가아웃 손실이 겹칠 때 경기 복구 루트가 잠기는 패턴이 더 뚜렷했다. 이 메시지를 후보 feature contract로 바꾼 뒤, 과거 KBO 외국인 성공/실패 classifier, KBO translation filter, market construction, failure risk gate, salary/contract gate, 그리고 네 명의 독립 후보 lens를 stacked-style 앙상블로 통합했다.",
        first_line=True,
    )
    add_p(
        doc,
        "순수 앙상블 보드에서 타자 Top 3는 Nolan Jones, Luis Matos, Jack Suwinski이며, 투수 raw Top 3는 Josh Fleming, Carson Spiers, Brian Van Belle이다. 그러나 최종 영입 의사결정은 모델 점수만으로 끝나지 않는다. 공개 연봉/계약/메디컬 gate를 붙이면 타자는 Nolan Jones가 모델 종합 1위이지만, Luis Matos가 가장 현실적인 최종 접촉 1순위로 올라온다. 투수는 Josh Fleming만 raw board와 salary/medical gate를 동시에 통과하는 1순위로 남고, Carson Spiers와 Brian Van Belle은 medical hold로 내려간다. 최종 결론은 타자 Luis Matos, 투수 Josh Fleming을 우선 접촉 대상으로 설정하되, Nolan Jones는 비용 조건이 맞을 경우 상향 가능한 model lead로 유지하는 것이다.",
        first_line=True,
    )
    add_heading(doc, "English", 2)
    add_p(
        doc,
        "This report reframes SSG Landers' replacement foreign-player search as an ensemble data-mining decision problem. The project first mines SSG-specific hidden needs, then converts them into candidate-level feature contracts and combines historical KBO foreign-player classifiers, KBO translation filters, market-access construction, failure-risk screens, salary/contract gates, and four independent analyst lenses. The raw hitter board favors Nolan Jones, Luis Matos, and Jack Suwinski, while the raw pitcher board begins with Josh Fleming, Carson Spiers, and Brian Van Belle. Once contract and medical gates are attached, Luis Matos becomes the most practical hitter contact and Josh Fleming remains the only pitcher lead that survives both model and feasibility screens.",
        first_line=True,
    )
    add_heading(doc, "Key Takeaways", 2)
    add_bullets(
        doc,
        [
            "SSG의 핵심 메시지는 ‘외야 장타 부족’이 아니라 ‘우투 선발전 OF/DH 전환 실패와 run-kill이 겹칠 때 경기 복구 루트가 잠기는 구조’다.",
            "네 명의 후보안은 단순 투표가 아니라 market lens, SSG fit lens, upside lens, supervised data-mining lens로 재정의했다.",
            "타자 raw Top 3는 Nolan Jones, Luis Matos, Jack Suwinski다. 최종 접촉 1순위는 salary gate를 통과하기 쉬운 Luis Matos다.",
            "투수 raw Top 3는 Josh Fleming, Carson Spiers, Brian Van Belle이나, Spiers와 Van Belle은 medical hold로 내려가며 Josh Fleming이 최종 1순위다.",
            "모델 결론은 ‘좋은 선수’가 아니라 ‘SSG 메시지, KBO 번역, 시장 접근성, 비용/메디컬 gate를 동시에 통과한 선수’를 찾는 절차다.",
        ],
    )
    start_page(doc)


def contents(doc: Document) -> None:
    add_heading(doc, "목차", 1)
    rows = [
        ("초록 (Abstract)", "i"),
        ("1 서론", "1"),
        ("  1.1 문제 정의", "1"),
        ("  1.2 핵심 지표", "1"),
        ("  1.3 분석 질문", "2"),
        ("2 방법", "3"),
        ("  2.1 데이터 및 전처리", "3"),
        ("  2.2 앙상블 base model 구성", "4"),
        ("  2.3 네 명의 분석 lens를 모델화하는 방식", "5"),
        ("  2.4 후보 점수화 및 gate 설계", "6"),
        ("3 결과", "7"),
        ("  3.1 SSG 숨은 약점 마이닝", "7"),
        ("  3.2 KBO 외인 성공/실패 모델", "9"),
        ("  3.3 후보 시장 구축과 번역 가능성", "10"),
        ("  3.4 타자 후보 Top 3와 최종 1인", "12"),
        ("  3.5 투수 후보 Top 3와 최종 1인", "14"),
        ("  3.6 앙상블 해석: 후보별 차별점", "16"),
        ("4 논의", "18"),
        ("5 한계", "20"),
        ("6 결론", "22"),
        ("A 부록", "24"),
        ("참고문헌", "26"),
    ]
    for left, right in rows:
        p = doc.add_paragraph()
        add_run(p, left, size=9.8)
        add_run(p, " " + "." * max(2, 82 - len(left)) + " ", size=9.8, color=MUTED)
        add_run(p, right, size=9.8)
    start_page(doc)


def list_pages(doc: Document) -> None:
    add_heading(doc, "표목차", 1)
    table_rows = [
        ("1.1", "주요 지표와 실무 해석", "1"),
        ("1.2", "분석 질문", "2"),
        ("2.1", "데이터 커버리지 요약", "3"),
        ("2.2", "base model 구성", "4"),
        ("2.3", "네 명의 분석 lens", "5"),
        ("2.4", "앙상블 weight", "6"),
        ("3.1", "SSG 숨은 약점 핵심 rule", "7"),
        ("3.2", "후보 feature contract", "8"),
        ("3.3", "모델 검증 요약", "9"),
        ("3.4", "타자 후보 Top 3", "12"),
        ("3.5", "투수 raw Top 3와 gate-pass 후보", "14"),
        ("6.1", "최종 의사결정표", "22"),
    ]
    for no, name, page in table_rows:
        p = doc.add_paragraph()
        add_run(p, f"{no}   {name} " + "." * 42 + f" {page}", size=9.4)
    add_heading(doc, "그림목차", 1)
    fig_rows = [
        ("3.1", "앙상블 signal weight", "6"),
        ("3.2", "타자 앙상블 랭킹", "12"),
        ("3.3", "투수 앙상블 랭킹", "14"),
        ("3.4", "모델 검증 AUC", "17"),
    ]
    for no, name, page in fig_rows:
        p = doc.add_paragraph()
        add_run(p, f"{no}   {name} " + "." * 52 + f" {page}", size=9.4)
    start_page(doc)


def intro(doc: Document) -> None:
    add_heading(doc, "1. 서론", 1)
    add_heading(doc, "1.1 문제 정의", 2)
    add_p(
        doc,
        "이 보고서는 SSG 랜더스가 외국인 타자 1명과 외국인 투수 1명을 영입한다고 가정할 때, 어떤 선수를 선택해야 하는지를 다룬다. 다만 문제의 단위는 선수 이름이 아니라 의사결정 규칙이다. 좋은 성적표를 가진 선수를 고르는 것이 아니라, SSG의 2026 경기 구조상 실제 약점을 보완하고 KBO에서 번역될 가능성이 있으며, 시장과 비용 조건까지 맞는 후보를 찾는 것이 목표다.",
        first_line=True,
    )
    add_p(
        doc,
        "따라서 후보 선정은 세 단계로 이루어진다. 첫째, SSG만의 약점을 상황별 interaction으로 찾는다. 둘째, 그 약점을 선수 feature로 바꾼다. 셋째, 네 명의 후보안을 base model lens로 통합한 뒤 salary/medical gate를 붙여 최종 접촉 우선순위를 정한다.",
        first_line=True,
    )
    add_heading(doc, "1.2 핵심 지표", 2)
    metrics = pd.DataFrame(
        [
            ["ensemble_score", "여러 base model 신호를 합성한 점수", "raw 후보 Top 3 산출"],
            ["dm_success_prob", "과거 KBO 외인 성공 label 기반 성공확률", "타자 후보 신뢰도"],
            ["dm_failure_prob", "과거 KBO 외인 실패 label 기반 실패위험", "하방 리스크"],
            ["consensus_signal", "서로 다른 분석 lens에서 반복 등장한 정도", "모델 안정성"],
            ["salary/medical gate", "연봉·소속·부상 상태 기반 hard gate", "최종 접촉 가능성"],
            ["RHP game-script lock", "상대 우투 선발전 SSG 공격/수비 손실 interaction", "SSG fit feature 정의"],
        ],
        columns=["지표", "의미", "사용 위치"],
    )
    add_table(doc, metrics, ["지표", "의미", "사용 위치"], ["지표", "의미", "사용 위치"], "표 1.1: 주요 지표와 실무 해석", size=8.1)
    add_heading(doc, "1.3 분석 질문", 2)
    rq = pd.DataFrame(
        [
            ["RQ1", "SSG의 숨은 약점은 무엇인가?", "후보 feature contract"],
            ["RQ2", "과거 KBO 외인 성공/실패 패턴에서 살아남는 후보는 누구인가?", "KBO translation 신뢰도"],
            ["RQ3", "네 명의 독립 후보안을 하나의 앙상블 모델로 어떻게 통합할 것인가?", "raw Top 3"],
            ["RQ4", "연봉/계약/메디컬 gate를 붙이면 순위가 어떻게 바뀌는가?", "실행 가능 Top 3"],
            ["RQ5", "최종 접촉 1순위 타자와 투수는 누구인가?", "최종 후보 1인"],
        ],
        columns=["질문", "내용", "산출물"],
    )
    add_table(doc, rq, ["질문", "내용", "산출물"], ["질문", "내용", "산출물"], "표 1.2: 분석 질문", size=8.1)


def method(doc: Document, t: dict[str, pd.DataFrame]) -> None:
    add_heading(doc, "2. 방법", 1)
    add_heading(doc, "2.1 데이터 및 전처리", 2)
    cov = t["coverage"].head(12).copy()
    cov["size_mb"] = cov["size_mb"].map(lambda x: f"{x:.1f}")
    cov["rows"] = cov["rows"].map(lambda x: f"{int(x):,}" if pd.notna(x) and x > 0 else "-")
    add_table(doc, cov, ["league_bucket", "files", "size_mb", "rows", "detail"], ["데이터", "파일", "MB", "rows", "설명"], "표 2.1: 데이터 커버리지 요약", size=6.8)
    add_p(
        doc,
        "KBO/STATIZ snapshot, MLB Savant pitch-level data, MiLB year-by-year official stats, MLB roster/transaction data, NPB official first-team/farm stats를 결합했다. 핵심은 후보 시장을 넓히는 것이 아니라, 각 후보가 실제로 어느 수준의 데이터 근거를 갖는지 표시하는 것이다.",
        first_line=True,
    )
    add_heading(doc, "2.2 앙상블 base model 구성", 2)
    add_table(doc, t["model_blocks"], ["model", "name", "data", "features", "role"], ["모델", "이름", "데이터", "주요 feature", "역할"], "표 2.2: base model 구성", size=6.8)
    add_heading(doc, "2.3 네 명의 분석 lens를 모델화하는 방식", 2)
    add_table(doc, t["team_lens"], ["analyst", "model_lens", "main_signal", "main_candidates", "ensemble_role"], ["분석", "모델 lens", "핵심 신호", "주요 후보", "앙상블 내 역할"], "표 2.3: 네 명의 분석 lens", size=6.8)
    add_p(
        doc,
        "여기서 앙상블은 단순 다수결이 아니다. 각자의 후보표를 하나의 예측기로 보고, 어떤 후보가 서로 다른 가정에서도 반복적으로 살아남는지를 측정한다. 예를 들어 Nolan Jones는 SSG fit lens와 supervised model에서 동시에 올라오고, Josh Fleming은 세 개 lens에서 반복 등장한다. 반대로 Jack Suwinski는 upside lens에서는 강하지만 supervised 성공확률에서는 약하다.",
        first_line=True,
    )
    add_heading(doc, "2.4 후보 점수화 및 gate 설계", 2)
    w = t["weights"].copy()
    w["weight"] = w["weight"].map(lambda x: f"{x:.0%}")
    add_table(doc, w, ["slot", "signal", "weight", "reason"], ["슬롯", "signal", "비중", "이유"], "표 2.4: 앙상블 weight", size=6.7)
    add_picture(doc, ASSETS / "ensemble_signal_weights.png", "그림 2.1: 슬롯별 앙상블 signal weight")
    add_callout(
        doc,
        "점수화 원칙.",
        "타자는 historical classifier 성능이 상대적으로 좋아 supervised signal 비중을 크게 둔다. 투수는 classifier 성능이 약하므로 consensus, KBO adaptation, market feature의 비중을 높인다.",
        fill=GOLD_LIGHT,
    )


def results(doc: Document, t: dict[str, pd.DataFrame]) -> None:
    add_heading(doc, "3. 결과", 1)
    add_heading(doc, "3.1 RQ1: SSG 숨은 약점 마이닝", 2)
    add_table(doc, t["hidden_rules"], ["rule", "interaction", "games", "win_pct", "run_diff", "interpretation"], ["rule", "interaction", "G", "승률", "득실", "해석"], "표 3.1: SSG 숨은 약점 핵심 rule", size=7.1)
    add_p(
        doc,
        "가장 중요한 인사이트는 SSG의 문제가 단순히 외야 장타 부족이 아니라는 점이다. OF HR zero만 보면 승률 .405, 평균 득실 -1.29 수준으로 너무 넓다. 반면 상대 우투 선발, OF/DH 전환 실패, run-kill, extra-out 손실이 결합하면 승률이 .100 또는 .000까지 떨어지고 평균 득실도 -4.50~-5.83으로 악화된다.",
        first_line=True,
    )
    add_table(
        doc,
        t["feature_blueprint"].head(8),
        ["slot", "message_component", "candidate_feature", "desired_direction", "measurement_proxy", "downstream_layer", "hard_gate_or_weight"],
        ["슬롯", "메시지", "선수 feature", "방향", "측정 proxy", "연결 layer", "가중/게이트"],
        "표 3.2: SSG 메시지를 선수 feature로 바꾸는 contract",
        size=5.9,
    )
    add_heading(doc, "3.2 RQ2: KBO 외인 성공/실패 모델", 2)
    audit = t["model_audit"].copy()
    add_table(doc, audit, ["slot", "model_block", "historical_rows", "feature_family", "model", "target", "feature_count", "validation_summary"], ["슬롯", "모델", "rows", "feature", "알고리즘", "target", "p", "검증 요약"], "표 3.3: 모델 검증 요약", size=6.2)
    add_picture(doc, ASSETS / "model_validation_auc.png", "그림 3.1: 타자/투수 classifier 검증 성능")
    add_p(
        doc,
        "타자 모델은 Savant 기반 ridge logistic classifier가 추천 모델로 쓸 수 있는 수준까지 올라왔다. 반면 투수 모델은 sparse L1 logistic을 사용했지만 검증 성능이 약해 diagnostic model로 둔다. 이 차이 때문에 타자 결론과 투수 결론의 표현 강도가 달라진다.",
        first_line=True,
    )
    add_heading(doc, "3.3 RQ3: 후보 시장 구축과 번역 가능성", 2)
    market = pd.DataFrame(
        [
            ["NPB 공식 2026", "2,186 rows", "1군/팜 타격·투구", "아쿼/아시아 시장 depth 파악"],
            ["MiLB current market", "18,258 rows", "1,745명 전수 AAA/AA 등급 조회", "현 후보의 최근 역할·레벨 확인"],
            ["Historical KBO pre-arrival MiLB", "71 feature rows", "KBO 외인 도착 전 MiLB 이력", "성공/실패 archetype 학습"],
            ["MLB transactions", "11,799 rows", "DFA, outright, minor deal, release", "시장 접근성 gate"],
        ],
        columns=["데이터", "규모", "내용", "후보 선정 역할"],
    )
    add_table(doc, market, ["데이터", "규모", "내용", "후보 선정 역할"], ["데이터", "규모", "내용", "후보 선정 역할"], "표 3.4: 후보 시장 구축 결과", size=7.4)
    add_heading(doc, "3.4 RQ4: 타자 후보 Top 3와 최종 1인", 2)
    add_picture(doc, ASSETS / "hitter_ensemble_ranking.png", "그림 3.2: 외국인 타자 앙상블 랭킹")
    add_table(doc, t["hitter_top"], ["rank", "candidate", "ensemble_score_fmt", "success_prob", "failure_prob", "source_presence", "ensemble_tier"], ["순위", "선수", "점수", "성공확률", "실패확률", "근거 lens", "판정"], "표 3.5: 타자 raw Top 3", size=7.3)
    add_table(doc, t["hitter_dm_view"], ["player_name", "roster_team", "age", "recent_pa", "recent_woba", "recent_bb_pct", "recent_k_pct", "recent_hardhit_rate", "success_prob", "failure_prob", "data_mining_rank"], ["선수", "팀", "나이", "PA", "wOBA", "BB%", "K%", "HardHit", "성공", "실패", "DM rank"], "표 3.6: 타자 data-mining 세부 지표", size=6.4)
    add_table(doc, t["salary_gate"][t["salary_gate"]["slot"].eq("foreign_hitter")], ["candidate", "salary_fmt", "contract_access_status", "medical_availability_status", "economic_gate", "model_action"], ["선수", "연봉 신호", "계약 상태", "가용성", "비용 gate", "조치"], "표 3.7: 타자 salary/contract gate", size=6.3)
    add_callout(
        doc,
        "타자 최종 결론.",
        "raw ensemble은 Nolan Jones를 1위로 본다. 그러나 비용 gate를 붙이면 Luis Matos가 최종 접촉 1순위다. Nolan Jones는 비용 조건이 정리될 때 즉시 상향 가능한 model lead로 유지한다.",
        fill=GREEN_LIGHT,
    )
    add_heading(doc, "3.5 RQ5: 투수 후보 Top 3와 최종 1인", 2)
    add_picture(doc, ASSETS / "pitcher_ensemble_ranking.png", "그림 3.3: 외국인 투수 앙상블 랭킹")
    add_table(doc, t["pitcher_raw"], ["rank", "candidate", "ensemble_score_fmt", "source_presence", "ensemble_tier"], ["순위", "선수", "점수", "근거 lens", "판정"], "표 3.8: 투수 raw Top 3", size=7.3)
    add_table(doc, t["pitcher_gate"], ["player_name", "roster_team", "age", "is_40man", "status_description", "market_access_bucket", "pre_kbo_milb_ip", "pre_kbo_milb_k9", "pre_kbo_milb_bb9", "success_prob", "failure_prob", "data_mining_rank"], ["선수", "팀", "나이", "40인", "상태", "시장", "MiLB IP", "K/9", "BB/9", "성공", "실패", "DM rank"], "표 3.9: 투수 gate-pass 후보 재정렬", size=5.9)
    add_table(doc, t["salary_gate"][t["salary_gate"]["slot"].eq("foreign_pitcher")], ["candidate", "salary_fmt", "contract_access_status", "medical_availability_status", "economic_gate", "model_action"], ["선수", "연봉 신호", "계약 상태", "가용성", "비용 gate", "조치"], "표 3.10: 투수 salary/medical gate", size=6.2)
    add_callout(
        doc,
        "투수 최종 결론.",
        "Josh Fleming은 raw board 1위이면서 계약 접근성도 가장 낫다. Carson Spiers와 Brian Van Belle은 모델상 후보였지만 medical gate에서 내려간다. Bryse Wilson/Austin Gomber는 gate-pass diagnostic 후보이나 consensus가 약해 최종 1인을 뒤집지 못한다.",
        fill=GREEN_LIGHT,
    )
    add_heading(doc, "3.6 앙상블 해석: 후보별 차별점", 2)
    differentiation = pd.DataFrame(
        [
            ["Luis Matos", "타자", "DM rank 1, 낮은 K%, 낮은 salary signal", "consensus 약함", "최종 접촉 1순위"],
            ["Nolan Jones", "타자", "앙상블 1위, 성공확률 90.2%, consensus 1.00", "2.0M 비용 gate", "조건부 최상위 model lead"],
            ["Jack Suwinski", "타자", "BB%, barrel/upside, 두 분석 lens에서 반복", "실패확률 높고 K% 큼", "upside hold"],
            ["Josh Fleming", "투수", "세 분석 lens 반복, minor contract 접근성", "투수 classifier 자체는 diagnostic", "최종 접촉 1순위"],
            ["Bryse Wilson", "투수", "gate-pass, non40man, 103 IP", "consensus 부족", "비교 검증 후보"],
            ["Austin Gomber", "투수", "gate-pass, 172 IP, 좌완 선발 이력", "성공확률 margin 작음", "비교 검증 후보"],
        ],
        columns=["선수", "슬롯", "강점", "리스크", "결론"],
    )
    add_table(doc, differentiation, ["선수", "슬롯", "강점", "리스크", "결론"], ["선수", "슬롯", "강점", "리스크", "결론"], "표 3.11: 후보별 차별점", size=7.3)


def discussion(doc: Document, t: dict[str, pd.DataFrame]) -> None:
    add_heading(doc, "4. 논의", 1)
    add_heading(doc, "4.1 실무 적용 원칙", 2)
    add_p(
        doc,
        "이번 분석의 가장 큰 실무적 의미는 후보 추천을 ‘성적 좋은 선수 고르기’에서 ‘SSG의 특정 경기 붕괴 패턴을 줄이는 선수 찾기’로 바꿨다는 점이다. 타자는 우투 선발전 OF/DH 이닝 복구, run-kill 회피, 두 스트라이크 생존력이 중요하고, 투수는 traffic command와 extra-out resilience가 중요하다.",
        first_line=True,
    )
    add_heading(doc, "4.2 모델별 해석 원칙", 2)
    add_bullets(
        doc,
        [
            "historical classifier는 타자에게 강하게 쓰고, 투수에게는 약한 diagnostic 신호로만 쓴다.",
            "team lens consensus는 투수에서 더 중요하다. 투수 supervised model이 약하기 때문이다.",
            "salary/medical gate는 모델 점수보다 나중에 붙는 hard gate다. 즉 점수가 높아도 계약/메디컬에서 막히면 최종 후보가 아니다.",
            "후보 1인은 raw model lead가 아니라 최종 접촉 우선순위로 정의한다.",
        ],
    )
    add_heading(doc, "4.3 전략적 메시지", 2)
    add_callout(
        doc,
        "발표 메시지.",
        "우리는 후보를 먼저 정하고 이유를 붙인 것이 아니라, SSG의 경기 복구 루트가 잠기는 상황을 먼저 찾고 그 상황을 줄이는 선수를 찾았다. 최종 추천이 Matos/Fleming으로 정리되는 이유는 이들이 모델 신호와 현실 gate를 동시에 비교적 잘 통과하기 때문이다.",
        fill=GOLD_LIGHT,
    )


def limitations_conclusion(doc: Document, t: dict[str, pd.DataFrame]) -> None:
    add_heading(doc, "5. 한계", 1)
    add_heading(doc, "5.1 데이터 범위의 한계", 2)
    add_bullets(
        doc,
        [
            "현재 SSG hidden-weakness snapshot은 2026-06-11 기준 산출물이므로 최신 경기 반영이 필요하다.",
            "투수 historical classifier의 AUC가 낮아 투수 후보는 확정 예측이 아니라 검증 우선순위로 표현해야 한다.",
            "공개 연봉은 remaining salary, buyout, opt-out, 이적료, 한국행 의사까지 완전히 대체하지 못한다.",
            "NPB/CPBL 쪽은 성적 커버리지는 확장됐지만 salary/contract/buyout 데이터가 아직 충분하지 않다.",
        ],
    )
    add_heading(doc, "5.2 모델링의 한계", 2)
    add_p(
        doc,
        "과거 KBO 외국인 표본은 구조적으로 작다. 따라서 복잡한 black-box 모델보다 ridge logistic, sparse L1 logistic, feature-family 비교, gate 기반 앙상블처럼 설명 가능한 모델을 우선했다. 이는 예측 성능 극대화보다 발표와 실무 의사결정의 재현성을 우선한 선택이다.",
        first_line=True,
    )
    add_heading(doc, "6. 결론", 1)
    add_heading(doc, "6.1 분석 질문별 요약 답변", 2)
    answers = pd.DataFrame(
        [
            ["RQ1", "숨은 약점", "우투 선발전 OF/DH 전환 실패와 run-kill/extra-out 손실이 결합될 때 경기 복구 루트가 잠긴다."],
            ["RQ2", "KBO 패턴", "타자는 classifier를 추천에 활용 가능하나, 투수는 diagnostic으로 제한해야 한다."],
            ["RQ3", "앙상블", "네 명의 후보안은 서로 다른 model lens로 통합했고, 반복 등장 후보에 안정성 신호를 부여했다."],
            ["RQ4", "gate 변화", "Nolan Jones는 model lead지만 비용 gate에서 보류, Luis Matos는 최종 접촉 1순위로 상승한다."],
            ["RQ5", "최종 1인", "타자 Luis Matos, 투수 Josh Fleming을 최종 접촉 1순위로 둔다."],
        ],
        columns=["RQ", "주제", "요약 답변"],
    )
    add_table(doc, answers, ["RQ", "주제", "요약 답변"], ["RQ", "주제", "요약 답변"], "표 6.1: 분석 질문별 요약 답변", size=7.5)
    add_heading(doc, "6.2 최종 후보 보드", 2)
    final_board = pd.DataFrame(
        [
            ["외국인 타자", "Luis Matos", "Nolan Jones", "Jack Suwinski", "Matos 최종 접촉 1순위, Nolan 조건부 model lead"],
            ["외국인 투수", "Josh Fleming", "Bryse Wilson", "Austin Gomber", "Fleming 최종 접촉 1순위, Wilson/Gomber 비교 검증"],
        ],
        columns=["슬롯", "1순위", "2순위", "3순위", "운영 결론"],
    )
    add_table(doc, final_board, ["슬롯", "1순위", "2순위", "3순위", "운영 결론"], ["슬롯", "1순위", "2순위", "3순위", "운영 결론"], "표 6.2: salary/medical gate 반영 최종 후보 보드", size=8.0)
    add_table(doc, t["final_decision"], ["slot", "board", "final_one", "decision", "reason"], ["슬롯", "보드", "최종 1인", "판정", "근거"], "표 6.3: raw board와 gated board의 차이", size=7.3)
    add_heading(doc, "6.3 핵심 한 줄", 2)
    add_callout(
        doc,
        "최종 결론.",
        "SSG는 외야 홈런 보강이 아니라 우투 선발전 경기 복구 루트를 열어줄 타자와 traffic-command형 투수를 찾아야 한다. 그 조건을 데이터 마이닝과 현실 gate로 통과한 최종 접촉 1순위는 타자 Luis Matos, 투수 Josh Fleming이다.",
        fill=GREEN_LIGHT,
    )


def appendix(doc: Document, t: dict[str, pd.DataFrame]) -> None:
    start_page(doc)
    add_heading(doc, "A. 부록", 1)
    add_heading(doc, "A.1 실패위험 gate 요약", 2)
    add_table(doc, t["layer5"], ["fit_slot", "failure_risk_band_v0_3", "manual_review_tier_v0_3", "locked_rows", "release_allowed"], ["슬롯", "risk band", "review tier", "rows", "release"], "표 A.1: failure risk gate 요약", size=7.0, max_rows=12)
    add_heading(doc, "A.2 fit ranking stage 요약", 2)
    add_table(doc, t["layer6"], ["fit_slot", "ranking_stage_gate_v0_3", "locked_rows", "release_allowed"], ["슬롯", "stage", "rows", "release"], "표 A.2: fit ranking stage 요약", size=7.0, max_rows=14)
    add_heading(doc, "참고문헌", 1)
    add_table(doc, t["references"], ["source", "use", "url"], ["출처", "사용 목적", "URL"], "표 R.1: 공개 근거 링크", size=6.3)
    add_heading(doc, "재현 산출물", 1)
    add_bullets(
        doc,
        [
            "outputs/tables/ensemble_candidate_scores_v1.csv",
            "outputs/tables/data_mining_hitter_candidates_v1.csv",
            "outputs/tables/data_mining_pitcher_candidates_v1.csv",
            "outputs/tables/final_candidate_salary_contract_gate_v1.csv",
            "outputs/tables/ssg_layer1_candidate_feature_blueprint_v4.csv",
            "scripts/build_reference_style_final_report_v4.py",
        ],
    )


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    tables = prep_tables()
    doc = Document()
    set_document_style(doc)
    cover(doc)
    abstract(doc)
    contents(doc)
    list_pages(doc)
    intro(doc)
    method(doc, tables)
    results(doc, tables)
    discussion(doc, tables)
    limitations_conclusion(doc, tables)
    appendix(doc, tables)
    add_footer(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
