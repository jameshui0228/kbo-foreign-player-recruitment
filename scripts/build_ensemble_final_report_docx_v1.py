#!/usr/bin/env python3
"""Create a polished DOCX version of the ensemble final report."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TABLE_DIR = PROJECT_ROOT / "outputs" / "tables"
REPORT_DIR = PROJECT_ROOT / "outputs" / "reports"
DOC_OUT = PROJECT_ROOT / "output" / "doc"
DOCX_PATH = DOC_OUT / "ssg_foreign_player_ensemble_data_mining_final_report_v1.docx"

SCORES_PATH = TABLE_DIR / "ensemble_candidate_scores_v1.csv"
WEIGHTS_PATH = TABLE_DIR / "ensemble_model_signal_weights_v1.csv"
ASSET_DIR = REPORT_DIR / "ensemble_final_report_v1_assets"


FONT = "Apple SD Gothic Neo"
INK = "1F2430"
MUTED = "6F768A"
GRID = "D7DBE7"
BLUE = "5477C4"
BLUE_LIGHT = "EAF1FE"
ORANGE_LIGHT = "FFEDDE"
GOLD_LIGHT = "FFF4C2"
NEUTRAL_LIGHT = "F4F5F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = GRID, size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 9.2) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 23, INK),
        ("Heading 1", 16, INK),
        ("Heading 2", 12.8, INK),
        ("Subtitle", 10.5, MUTED),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if style_name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)


def add_run(paragraph, text: str, *, bold: bool = False, color: str = INK, size: float | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_paragraph(doc: Document, text: str = "", *, bold_prefix: str | None = None, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(paragraph, bold_prefix, bold=True)
        add_run(paragraph, text[len(bold_prefix):])
    else:
        add_run(paragraph, text)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    add_run(paragraph, text, bold=True, size=16 if level == 1 else 12.8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_run(paragraph, item)


def add_metric_cards(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    metrics = [
        ("0.833", "타자 성공 classifier AUC"),
        ("0.738", "타자 실패 classifier AUC"),
        ("0.603", "투수 성공 classifier AUC, watch only"),
    ]
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, NEUTRAL_LIGHT)
        set_cell_border(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(paragraph, value, bold=True, size=18)
        paragraph.add_run("\n")
        add_run(paragraph, label, color=MUTED, size=9)


def add_table(doc: Document, data: pd.DataFrame, columns: list[str], headers: list[str], *, max_rows: int | None = None) -> None:
    if max_rows is not None:
        data = data.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, header in zip(hdr, headers):
        set_cell_shading(cell, NEUTRAL_LIGHT)
        set_cell_border(cell)
        set_cell_text(cell, header, bold=True, size=8.5)
    for _, row in data.iterrows():
        cells = table.add_row().cells
        for cell, col in zip(cells, columns):
            value = row[col]
            if pd.isna(value):
                text = "-"
            elif col in {"ensemble_score", "dm_success_prob", "dm_failure_prob", "dm_margin", "dm_signal", "consensus_signal"}:
                text = f"{float(value) * 100:.1f}%"
            else:
                text = str(value)
            set_cell_border(cell)
            set_cell_text(cell, text, size=8.2)


def add_picture(doc: Document, path: Path, caption: str, *, width: float = 6.7) -> None:
    if not path.exists():
        add_paragraph(doc, f"[Missing chart: {path.name}]")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(caption_p, caption, color=MUTED, size=8.7)


def add_callout(doc: Document, title: str, body: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=GRID)
    paragraph = cell.paragraphs[0]
    add_run(paragraph, title, bold=True)
    add_run(paragraph, " " + body)


def pct_or_dash(value) -> str:
    if pd.isna(value):
        return "-"
    return f"{float(value) * 100:.1f}%"


def build_docx() -> None:
    scores = pd.read_csv(SCORES_PATH)
    weights = pd.read_csv(WEIGHTS_PATH)
    DOC_OUT.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_document_defaults(doc)

    # Title page
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title, "SSG 외국인 선수 영입 앙상블\n데이터 마이닝 최종 보고서", bold=True, size=23)
    subtitle = doc.add_paragraph(style="Subtitle")
    add_run(subtitle, "SDA 2기 SSG 대체 외국인/아시아쿼터 선수 영입 프로젝트 | 2026-06-23", color=MUTED, size=10.5)

    add_callout(
        doc,
        "기술 요약.",
        "네 명의 분석을 따로 병렬 비교하지 않고, 하나의 stacked-style 앙상블 모델로 재구성했다. base learner는 과거 KBO 외국인 성공/실패 학습 모델, SSG fit/번역 파이프라인, KBO 적응 filter, 시장 비효율 feature model, 반복 등장 consensus signal이다.",
        BLUE_LIGHT,
    )
    add_callout(
        doc,
        "최종 결론.",
        "타자는 Nolan Jones가 앙상블 1순위이고, Luis Matos는 순수 성공/실패 classifier가 새로 끌어올린 model-discovery 후보로 남는다. 투수는 Josh Fleming만 consensus가 강하며, 나머지는 확정 추천이 아니라 검증 board로 두는 것이 타당하다.",
        GOLD_LIGHT,
    )
    add_metric_cards(doc)

    add_heading(doc, "1. 결론은 팀원별 후보 모음이 아니라 하나의 앙상블 모델이다", 1)
    add_paragraph(
        doc,
        "이번 최종 모델은 네 명의 산출물을 하나의 evidence stack으로 묶는다. Codex의 과거 KBO 성공/실패 classifier는 historical label을 직접 학습한 가장 강한 타자 base learner다. sewon의 6단계 파이프라인은 SSG 약점, KBO 번역, 영입 현실성을 연결한다. jimini의 filter는 KBO 변화구, 저속 구종, ABS 환경에서 무너질 후보를 제거한다. kyuho의 feature model은 MLB fringe 시장에서 KBO로 번역될 수 있는 숨은 장점을 찾는다.",
    )
    add_paragraph(
        doc,
        "따라서 발표에서는 각자 후보가 다르다는 식이 아니라, 서로 다른 base learner가 같은 후보를 얼마나 지지하는가를 보여주면 된다. 이 구조가 데이터 마이닝의 느낌을 가장 잘 살린다.",
    )
    add_picture(
        doc,
        ASSET_DIR / "ensemble_signal_weights.png",
        "그림 1. 슬롯별 base learner 비중. 타자는 classifier 검증력이 있어 historical model 비중이 크고, 투수는 consensus와 filter 비중이 커진다.",
    )

    add_heading(doc, "2. 모델 설계: base learner와 역할", 1)
    add_paragraph(
        doc,
        "앙상블은 단순 평균이 아니다. 각 base learner의 신뢰도와 역할이 다르기 때문에 슬롯별로 다른 weight를 부여했다. 타자는 과거 KBO 성공/실패 classifier가 검증을 통과했으므로 40%를 반영했고, 투수는 classifier AUC가 0.603에 그쳐 5% diagnostic signal로만 반영했다.",
    )
    weight_table = weights.copy()
    weight_table["weight"] = weight_table["weight"].map(lambda v: f"{v:.0%}")
    add_table(
        doc,
        weight_table,
        ["slot", "signal", "weight", "reason"],
        ["슬롯", "base learner", "비중", "근거"],
    )

    add_heading(doc, "3. 타자: Nolan Jones는 모델 확률과 consensus가 동시에 붙는다", 1)
    add_paragraph(
        doc,
        "타자 앙상블의 핵심은 Nolan Jones와 Luis Matos의 역할 차이다. Nolan Jones는 과거 성공/실패 classifier에서 높은 성공확률을 받고 sewon, jimini, Codex 구조화 모델에서도 반복 등장한다. 즉 모델 확률과 독립 분석이 동시에 지지한다. Luis Matos는 팀원 교집합은 약하지만 classifier가 가장 강하게 끌어올린 hidden/model-discovery 후보다.",
    )
    add_picture(
        doc,
        ASSET_DIR / "hitter_ensemble_ranking.png",
        "그림 2. 타자 상위 후보. Nolan Jones는 classifier와 consensus가 결합되어 1위가 되었고, Luis Matos는 classifier가 새로 발굴한 후보로 유지된다.",
    )
    hitter_top = scores[scores["slot"].eq("foreign_hitter")].head(8)
    add_table(
        doc,
        hitter_top,
        ["rank", "candidate", "ensemble_score", "dm_success_prob", "dm_failure_prob", "source_presence", "ensemble_tier"],
        ["순위", "선수", "앙상블 점수", "성공확률", "실패확률", "근거 출처", "판정"],
    )

    add_heading(doc, "4. 투수: Josh Fleming은 합의 후보지만 아직 확정 추천은 아니다", 1)
    add_paragraph(
        doc,
        "투수는 타자와 같은 방식으로 확정 추천하면 안 된다. 과거 KBO 투수 성공 classifier의 AUC는 0.603으로 watch 등급이다. 따라서 투수 모델은 누가 확실히 성공한다가 아니라 어떤 후보를 먼저 검증할지 정하는 장치다.",
    )
    add_paragraph(
        doc,
        "이 조건에서 Josh Fleming이 가장 앞서는 이유는 명확하다. jimini와 kyuho가 동시에 지지했고, 기존 Codex 통합모델에서도 반복 등장했다. 반면 Bryse Wilson과 Austin Gomber는 pure diagnostic model에서는 잡히지만 팀원 consensus가 약하다. Carson Spiers와 Brian Van Belle은 sewon의 번역/이닝형 후보로 가치가 있지만 medical, 계약 접근성, 역할 지속성 확인이 필요하다.",
    )
    add_picture(
        doc,
        ASSET_DIR / "pitcher_ensemble_ranking.png",
        "그림 3. 투수 상위 후보. Josh Fleming은 반복 등장 consensus가 강하지만, 투수 classifier 검증력이 낮아 최종 추천이 아니라 1차 검증 후보로 제시한다.",
    )
    pitcher_top = scores[scores["slot"].eq("foreign_pitcher")].head(8)
    add_table(
        doc,
        pitcher_top,
        ["rank", "candidate", "ensemble_score", "dm_success_prob", "dm_failure_prob", "source_presence", "ensemble_tier"],
        ["순위", "선수", "앙상블 점수", "성공확률", "실패위험", "근거 출처", "판정"],
    )

    add_heading(doc, "5. 왜 타자와 투수를 다르게 말해야 하는가", 1)
    add_paragraph(
        doc,
        "모델 검증 결과가 결론의 말투를 결정한다. 타자 성공/실패 classifier는 각각 AUC 0.833, 0.738로 후보 추천에 쓸 수 있는 수준이다. 반면 투수 성공 classifier는 AUC 0.603이라서 확정 추천으로 쓰기 어렵다. 이 차이를 숨기면 발표가 약해진다. 오히려 이 차이를 명확히 말해야 모델이 정직해 보인다.",
    )
    add_picture(
        doc,
        ASSET_DIR / "model_validation_auc.png",
        "그림 4. 검증 성능 비교. 타자는 promoted, 투수는 watch 등급으로 처리했기 때문에 최종 추천 강도도 달라진다.",
    )

    add_heading(doc, "6. 발표용 핵심 인사이트", 1)
    add_bullets(
        doc,
        [
            "SSG의 타자 보강은 장타 총량이 아니라 game-state conversion 문제다. 1루 주자 상황, 2사, 초반 이닝에서 공격 흐름을 살릴 수 있는 선구/갭파워/구종 대응형 외야 자원이 필요하다.",
            "외국인 타자 모델은 후보를 실제로 바꿨다. 기존 수동 fit 점수에서는 덜 보였던 Luis Matos가 classifier에서 강하게 올라왔고, Nolan Jones는 classifier와 consensus를 동시에 통과했다.",
            "투수는 정답 후보보다 검증 우선순위가 결론이다. Josh Fleming이 가장 강한 1차 검증 후보지만, 투수 전체는 계약, medical, 선발 지속성, BB/9와 HR/9 추가 확인이 필요하다.",
            "Dylan Carlson, Bryse Wilson, Austin Gomber는 model-only 또는 diagnostic 후보로 분리해야 한다. 이들을 무리하게 최종 추천처럼 말하지 않는 것이 보고서 신뢰도를 높인다.",
            "Jack Suwinski와 Michael Toglia는 최종 모델의 반박 후보군이다. classifier top은 아니지만 시장 비효율과 power-upside base learner가 지지하기 때문에 스카우팅 검증군에 남길 가치가 있다.",
        ],
    )

    add_heading(doc, "7. 최종 후보 board", 1)
    add_callout(
        doc,
        "최종 타자 board.",
        "Nolan Jones를 ensemble 1순위, Luis Matos를 model-discovery 1순위, Jack Suwinski와 Michael Toglia를 upside 반박 후보로 둔다. Dylan Carlson은 classifier Top 3지만 다른 base learner 지지가 약해 보류성 후보로 설명한다.",
        BLUE_LIGHT,
    )
    add_callout(
        doc,
        "최종 투수 board.",
        "Josh Fleming을 consensus 검증 1순위로 두고, Carson Spiers, Brian Van Belle, Bryse Wilson, Austin Gomber를 추가 검증군으로 둔다. 투수는 확정 추천이 아니라 검증 우선순위라는 표현을 유지한다.",
        ORANGE_LIGHT,
    )

    add_heading(doc, "8. 한계와 다음 검증", 1)
    add_paragraph(
        doc,
        "이 보고서는 현재 확보된 구조화 숫자 데이터로 만든 앙상블이다. 기사, 인터뷰, 평판 텍스트는 input에서 제외했다. 남은 핵심 리스크는 실제 연봉, buyout, opt-out, 한국행 의사, medical, 40인/계약 변동이다. 특히 투수는 이 추가 정보가 들어오기 전까지 추천보다 검증 board로 두는 편이 맞다.",
    )
    add_bullets(
        doc,
        [
            "타자 추가 검증: 2026 최신 Savant 표본, 좌/우 split, breaking/off-speed 대응, KBO 외야 수비/주루 역할.",
            "투수 추가 검증: 최근 4주 구속 변화, starter pitch count, BB/9와 HR/9의 park-adjusted 안정성, medical flag.",
            "계약 검증: 실제 연봉, 잔여 보장액, 방출 가능성, KBO 외국인 선수 계약 상한/세금/이적료 리스크.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "SDA 2기 SSG 외국인 선수 영입 프로젝트 | Ensemble Data Mining Report", color=MUTED, size=8)

    doc.save(DOCX_PATH)


def main() -> None:
    build_docx()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
