#!/usr/bin/env python3
"""Build an academic-style report DOCX matching the reference SDA report format."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "tables"
DOC_DIR = ROOT / "output" / "doc"
ASSET_DIR = ROOT / "tmp" / "docs" / "academic_report_v6_assets"
DOC_PATH = DOC_DIR / "SSG_외국인선수_데이터마이닝_학술보고서_보강_v6.docx"

REPORT_TITLE = "SSG 랜더스 외국인 선수 영입 데이터 마이닝 전략"
REPORT_SUBTITLE = "구조화 데이터 기반 팀-선수 적합성 분석 보고서"
REPORT_DATE = "June 23, 2026"


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(OUT / name)


def fmt_num(value, digits: int = 3) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{float(value):.{digits}f}"


def fmt_int(value) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{int(round(float(value))):,}"


def fmt_pct(value, digits: int = 1) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{float(value) * 100:.{digits}f}%"


def set_run_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleMyungjo")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(document: Document, text: str = "", *, size: float = 10.5, align=None, first_line: bool = False) -> None:
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.45)
    run = p.add_run(text)
    set_run_font(run, size=size)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = p.add_run(text)
    set_run_font(run, size=15.5 if level == 1 else 13.2, bold=True)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=9.5)


def add_toc_line(document: Document, label: str, page: str, *, indent: int = 0) -> None:
    dots = "." * max(10, 70 - len(label) - indent * 4)
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55 * indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{label} {dots} {page}")
    set_run_font(run, size=10.5)


def set_cell_text(cell, text, *, size: float = 8.8, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run("" if text is None else str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        attrs = kwargs[edge]
        if attrs is None:
            attrs = {"val": "nil"}
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def add_academic_table(
    document: Document,
    caption: str,
    headers: list[str],
    rows: Iterable[Iterable[object]],
    *,
    widths: list[float] | None = None,
    size: float = 8.5,
) -> None:
    add_caption(document, caption)
    rows = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, size=size, bold=False)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, size=size, align=align)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell, left=None, right=None, top=None, bottom=None)
            if row_idx == 0:
                set_cell_border(cell, top={"val": "single", "sz": "10", "color": "000000"})
                set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": "000000"})
            if row_idx == len(table.rows) - 1:
                set_cell_border(cell, bottom={"val": "single", "sz": "10", "color": "000000"})
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)
    document.add_paragraph()


def add_key_takeaways(document: Document, bullets: list[str]) -> None:
    table = document.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.cell(0, 0)
    body = table.cell(1, 0)
    set_cell_text(header, "Key Takeaways", size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(body, "", size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for cell, fill in [(header, "BFBFBF"), (body, "F3F3F3")]:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    body.text = ""
    for bullet in bullets:
        p = body.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("• " + bullet)
        set_run_font(run, size=9.5)
    document.add_paragraph()


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9.5)


def set_page_numbering(section, *, start: int = 1, fmt: str = "decimal") -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def set_header_footer(section, *, title: str = REPORT_TITLE, date: str = REPORT_DATE, header_on: bool = True, page_number: bool = True) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    if header_on:
        p = section.header.paragraphs[0]
        p.text = ""
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(title)
        set_run_font(run, size=8.8)
        run = p.add_run("\t" + date)
        set_run_font(run, size=8.8)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
    else:
        section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""
    if page_number:
        add_page_number(section.footer.paragraphs[0])


def make_charts() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.family"] = ["AppleMyungjo", "AppleGothic", "Times New Roman", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    charts: dict[str, Path] = {}

    runway = read_csv("ssg_2026_runway_gap_by_team.csv")
    ssg = runway[runway["t_code_name"].eq("SSG")].iloc[0]
    fig, ax = plt.subplots(figsize=(5.6, 3.1))
    labels = ["RISP OPS\n(rank 1)", "1B OPS\n(rank 10)", "1B OBP\n(rank 10)"]
    values = [ssg["OPS_risp"], ssg["OPS_on_first"], ssg["OBP_on_first"]]
    ax.bar(labels, values, color=["#2F5597", "#7F7F7F", "#BFBFBF"], edgecolor="black", linewidth=0.5)
    ax.set_ylim(0, 0.95)
    ax.set_ylabel("Rate")
    for i, v in enumerate(values):
        ax.text(i, v + 0.025, f"{v:.3f}", ha="center", fontsize=9)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    path = ASSET_DIR / "figure_3_1_runway.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    charts["runway"] = path

    context = read_csv("ssg_pitching_message_v0_2_context_validation.csv").head(6)
    fig, ax = plt.subplots(figsize=(5.6, 3.1))
    ax.barh(context["context_label"][::-1], context["bad_context_score"][::-1], color="#7F7F7F", edgecolor="black", linewidth=0.5)
    ax.set_xlabel("Bad-context score")
    ax.set_xlim(0, 10.5)
    for i, v in enumerate(context["bad_context_score"][::-1]):
        ax.text(v + 0.12, i, f"{v:.1f}", va="center", fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    path = ASSET_DIR / "figure_3_2_pitching_context.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    charts["pitch_context"] = path

    top = read_csv("data_mining_recommendations_top3_v1.csv")
    fig, ax = plt.subplots(figsize=(5.8, 3.5))
    names = top["player_name"].tolist()[::-1]
    y = list(range(len(names)))
    ax.barh(y, top["dm_success_prob"].tolist()[::-1], color="#2F5597", edgecolor="black", linewidth=0.4, label="success")
    ax.barh(y, [-x for x in top["dm_failure_prob"].tolist()[::-1]], color="#BFBFBF", edgecolor="black", linewidth=0.4, label="failure")
    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_yticks(y)
    ax.set_yticklabels(names)
    ax.set_xlim(-0.65, 1.0)
    ax.set_xlabel("Probability")
    ax.legend(loc="lower right", fontsize=8, frameon=False)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    path = ASSET_DIR / "figure_3_3_candidates.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    charts["candidates"] = path

    fig, ax = plt.subplots(figsize=(7.2, 2.6))
    ax.axis("off")
    steps = [
        ("Input", "KBO splits\nSavant/MiLB\nroster/cost"),
        ("1", "SSG hidden\nweakness"),
        ("2", "KBO success/\nfailure type"),
        ("3", "Candidate\nmarket"),
        ("4-5", "Translation +\nrisk gates"),
        ("Output", "Hitter 3\nPitcher leads"),
    ]
    xs = [0.06, 0.23, 0.40, 0.57, 0.74, 0.91]
    for i, ((title, body), x) in enumerate(zip(steps, xs)):
        ax.text(
            x,
            0.58,
            f"{title}\n{body}",
            ha="center",
            va="center",
            fontsize=8.4,
            bbox=dict(boxstyle="round,pad=0.35", facecolor="#F2F2F2", edgecolor="#333333", linewidth=0.8),
        )
        if i < len(xs) - 1:
            ax.annotate(
                "",
                xy=(xs[i + 1] - 0.07, 0.58),
                xytext=(x + 0.07, 0.58),
                arrowprops=dict(arrowstyle="->", color="#333333", linewidth=1.0),
            )
    ax.text(
        0.5,
        0.12,
        "Decision rule: release candidates only after model pass + market gate + contract/medical/Korea-willingness review",
        ha="center",
        va="center",
        fontsize=8.2,
    )
    fig.tight_layout()
    path = ASSET_DIR / "figure_2_1_pipeline.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    charts["pipeline"] = path
    return charts


def setup_document(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2.25)
        section.bottom_margin = Cm(2.1)
        section.left_margin = Cm(2.55)
        section.right_margin = Cm(2.55)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "AppleMyungjo")
    styles["Normal"].font.size = Pt(10.5)


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(185)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_TITLE)
    set_run_font(run, size=23, bold=True)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_SUBTITLE)
    set_run_font(run, size=14)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(235)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Data: KBO/STATIZ 2023-2026  |  MLB Savant 2023-2026  |  Candidate-market 1,745+")
    set_run_font(run, size=11)
    for text, size in [
        ("SDA 2기 SSG 랜더스 프로젝트 팀", 11),
        ("고려대학교 SDA", 11),
        (REPORT_DATE, 11),
    ]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=size)


def add_abstract(document: Document) -> None:
    add_heading(document, "초록 (Abstract)", 1)
    add_heading(document, "국문", 2)
    ko = (
        "본 보고서는 SSG 랜더스의 대체 외국인 선수 영입 문제를 단순 선수 평가가 아니라 팀 상태와 후보 상태의 "
        "적합성 문제로 재정의한다. 분석은 기사, 인터뷰, 뉴스 텍스트를 최종 모델 입력에서 제외하고, KBO/STATIZ "
        "상황별 데이터, MLB Savant, MiLB 성적, 로스터 상태, 후보 시장 접근성, KBO 외국인 선수 성공/실패 라벨 등 "
        "숫자형 및 구조화 데이터만 사용했다. 1단계에서는 SSG의 숨은 약점을 탐색했고, 타격에서는 득점권 생산성이 "
        "아니라 1루 주자 상황의 전환 실패가, 투수에서는 외국인 선발 슬롯의 run prevention 및 traffic command 실패가 "
        "핵심 병목으로 확인됐다. 2-5단계에서는 KBO 외인 성공/실패 archetype, 후보 시장, KBO 번역 모델, 실패 리스크 "
        "모델을 구축했다. 최종적으로 타자 모델은 Savant 기반 성공/실패 classifier가 pilot-promote 수준의 검증 성능을 "
        "보였고, Luis Matos, Nolan Jones, Dylan Carlson을 우선 검토 후보로 제시한다. 반면 투수 모델은 watch 수준에 "
        "머물러 Bryse Wilson과 Austin Gomber를 확정 추천이 아닌 진단 리드로 제시한다."
    )
    add_paragraph(document, ko, first_line=True)
    add_heading(document, "English", 2)
    en = (
        "This report reframes SSG Landers' foreign-player replacement problem as a team-player fit problem rather than "
        "a generic talent search. The final model excludes articles, interviews, and text-derived signals, and relies only "
        "on numeric or structured data: KBO/STATIZ situational splits, MLB Savant features, MiLB performance, roster status, "
        "market access, and historical KBO foreign-player success/failure labels. The hidden-need layer identifies a hitter-side "
        "first-base runway bottleneck and a pitcher-side import-slot traffic-command failure. The subsequent layers mine KBO "
        "success archetypes, build the candidate market, estimate KBO translation, and screen failure risk. The hitter classifier "
        "is strong enough for pilot use and points to Luis Matos, Nolan Jones, and Dylan Carlson. The pitcher classifier remains "
        "diagnostic, so Bryse Wilson and Austin Gomber should be treated as verification leads rather than final recommendations."
    )
    add_paragraph(document, en)
    add_key_takeaways(
        document,
        [
            "문제: SSG의 외인 보강은 단순 장타/구속 보강이 아니라 팀 고유 game-state 병목을 해결하는 문제다.",
            "타자 메시지: 득점권이 아니라 1루 주자 상황을 살리는 first-base traffic converter가 필요하다.",
            "투수 메시지: 외국인 선발 슬롯은 load-bearing traffic-command starter로 재설계되어야 한다.",
            "후보 결론: 타자는 Matos/Jones/Carlson, 투수는 Wilson/Gomber를 진단 리드로 둔다.",
        ],
    )


def add_executive_summary(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "발표용 한 장 요약", 1)
    top = read_csv("data_mining_recommendations_top3_v1.csv")
    hitters = ", ".join(top[top["slot"].eq("foreign_hitter")]["player_name"].tolist())
    pitchers = ", ".join(top[top["slot"].eq("foreign_pitcher")]["player_name"].tolist()[:2])
    rows = [
        ["핵심 주장", "SSG의 보강 문제는 '외야 장타' 또는 '삼진형 투수'가 아니라 특정 game-state 병목을 줄이는 선수 유형 탐색이다."],
        ["타자 결론", "RISP OPS는 강하지만 1루 주자 상황 OPS/OBP가 낮다. 따라서 필요한 유형은 first-base traffic converter다."],
        ["투수 결론", "외국인 선발 슬롯이 국내 선발보다 안정적이지 못했다. 필요한 유형은 6이닝을 버티는 traffic-command starter다."],
        ["후보 출력", f"외인타자 보드: {hitters}. 외인투수는 확정 추천이 아니라 진단 리드: {pitchers}."],
        ["확정 전 gate", "salary, opt-out, buyout, assignment status, medical, Korea-willingness를 통과해야 실제 영입 후보가 된다."],
    ]
    add_academic_table(document, "표 ES.1: 최종 발표 요약", ["항목", "내용"], rows, widths=[2.8, 10.5], size=8.4)
    add_paragraph(
        document,
        "발표에서 가장 중요한 문장은 다음과 같다. SSG는 강한 득점권 결과를 더 강하게 만드는 선수가 아니라, "
        "득점권 이전 단계에서 공격 흐름을 끊기지 않게 만드는 외국인 타자를 찾아야 한다. 투수 쪽은 upside보다 "
        "외국인 슬롯의 최소 안정성을 회복하는 것이 먼저다.",
        first_line=True,
    )
    decision_rows = [
        ["Model-supported", "Matos, Jones, Carlson", "타자 Savant 성공/실패 classifier가 pilot-promote"],
        ["Diagnostic only", "Wilson, Gomber", "투수 모델은 watch 수준이라 영상/계약/메디컬 검증 후 판단"],
        ["Hold", "Asian quota", "국적, 아시아리그 이력, 200k USD cap 미통과 상태에서 이름 확정 금지"],
    ]
    add_academic_table(document, "표 ES.2: 발표 시 표현 강도", ["표현 강도", "대상", "이유"], decision_rows, widths=[3, 4, 6.5], size=8.2)


def add_contents(document: Document) -> None:
    add_heading(document, "목차", 1)
    for label, page, indent in [
        ("초록 (Abstract)", "i", 0),
        ("발표용 한 장 요약", "iii", 0),
        ("1 서론", "1", 0),
        ("1.1 문제 정의", "1", 1),
        ("1.2 핵심 지표", "1", 1),
        ("1.3 분석 질문", "2", 1),
        ("2 방법", "3", 0),
        ("2.1 데이터 및 전처리", "3", 1),
        ("2.2 6단계 데이터 마이닝 프레임", "4", 1),
        ("2.3 통합 모델 구조", "5", 1),
        ("2.4 모델 검증과 해석 경계", "5", 1),
        ("3 결과", "6", 0),
        ("3.1 SSG 숨은 약점: 타격 runway", "6", 1),
        ("3.2 SSG 숨은 약점: 투수 import-slot inversion", "7", 1),
        ("3.3 KBO 외국인 성공/실패 유형", "8", 1),
        ("3.4 후보 시장과 현실성 gate", "9", 1),
        ("3.5 최종 후보와 데이터 마이닝 통과 논리", "10", 1),
        ("3.6 탈락 후보 audit", "11", 1),
        ("3.7 후보별 통과 조건과 반증 조건", "12", 1),
        ("3.8 반박 가능성 점검", "13", 1),
        ("4 논의", "14", 0),
        ("4.5 민감도 및 강건성 점검", "15", 1),
        ("4.6 최종 의사결정 gate", "16", 1),
        ("5 한계", "17", 0),
        ("6 결론", "18", 0),
        ("부록", "19", 0),
    ]:
        add_toc_line(document, label, page, indent=indent)
    document.add_page_break()

    add_heading(document, "표목차", 1)
    table_entries = [
        ("1.1 주요 야구 지표 요약", "1"),
        ("1.2 분석 질문", "2"),
        ("2.1 데이터 구성 요약", "3"),
        ("2.2 6단계 분석 프레임", "4"),
        ("2.3 통합 모델 입출력", "5"),
        ("2.4 모델 검증 요약", "5"),
        ("3.1 SSG 상황별 runway gap", "6"),
        ("3.2 SSG 투수 worst context", "7"),
        ("3.3 KBO 외국인 archetype", "8"),
        ("3.4 후보 시장 구성", "9"),
        ("3.5 최종 외인타자 후보", "10"),
        ("3.6 외인투수 진단 리드", "11"),
        ("3.7 주요 탈락 후보 audit", "11"),
        ("3.8 후보별 통과 및 반증 조건", "12"),
        ("3.9 주요 반박과 대응 근거", "13"),
        ("4.2 공개 연봉 및 계약 현실성 gate", "14"),
        ("4.3 아시아쿼터 잠금 사유", "14"),
        ("4.4 민감도 및 강건성 점검", "15"),
        ("4.5 최종 의사결정 gate", "16"),
    ]
    for label, page in table_entries:
        add_toc_line(document, label, page, indent=0)
    document.add_page_break()

    add_heading(document, "그림목차", 1)
    for label, page in [
        ("2.1 통합 데이터 마이닝 파이프라인", "5"),
        ("3.1 SSG 타격 runway gap", "6"),
        ("3.2 SSG 투수 bad-context score", "7"),
        ("3.3 후보별 success/failure probability", "10"),
    ]:
        add_toc_line(document, label, page, indent=0)


def add_intro(document: Document) -> None:
    add_heading(document, "1. 서론", 1)
    add_heading(document, "1.1 문제 정의", 2)
    add_paragraph(
        document,
        "이 보고서는 외국인 선수 영입을 '좋은 선수 찾기'가 아니라 'SSG의 특정 game-state 결함을 줄이는 선수 유형 찾기'로 정의한다. "
        "따라서 후보 평가는 포지션, 장타, 구속 같은 단일 능력보다 SSG의 약점과 KBO 외인 성공/실패 패턴을 동시에 통과하는지를 기준으로 한다.",
        first_line=True,
    )
    add_heading(document, "1.2 핵심 지표", 2)
    metric_rows = [
        ["RISP OPS", "득점권 공격 생산성", "SSG가 이미 강한 영역인지 확인"],
        ["Runner-on-1B OPS/OBP", "1루 주자 상황 전환력", "타자 보강의 숨은 핵심 병목"],
        ["GDP/PA", "run-killing out 위험", "first-base traffic converter의 반대 신호"],
        ["WHIP/BB9/HR9", "출루 허용과 장타 허용", "traffic-command starter 검증"],
        ["IP/start", "선발 runway", "불펜 tax를 줄이는 직접 지표"],
        ["KBO success/failure label", "과거 외인 결과", "후보 번역 모델의 학습 target"],
    ]
    add_academic_table(document, "표 1.1: 주요 야구 지표 요약", ["지표", "의미", "실무 해석"], metric_rows, widths=[3, 4, 6])
    add_heading(document, "1.3 분석 질문", 2)
    add_paragraph(
        document,
        "본 보고서는 세부 분석 단계를 여섯 개로 나누되, 발표와 의사결정의 중심 질문은 아래 두 개로 압축한다. "
        "첫 번째 질문은 SSG가 어떤 유형을 필요로 하는지, 두 번째 질문은 그 유형을 기준으로 실제 후보가 누구인지에 답한다.",
        first_line=True,
    )
    rq_rows = [
        ["RQ1", "SSG의 숨은 보강 포인트는 무엇이며, 이를 어떤 외국인 선수 유형으로 번역할 수 있는가?"],
        ["RQ2", "그 유형을 기준으로 KBO 성공/실패, 후보 시장, 번역/리스크 모델을 통과한 최종 후보는 누구인가?"],
    ]
    add_academic_table(document, "표 1.2: 분석 질문", ["질문", "의미"], rq_rows, widths=[2, 11.5])


def add_methods(document: Document, charts: dict[str, Path]) -> None:
    add_heading(document, "2. 방법", 1)
    add_heading(document, "2.1 데이터 및 전처리", 2)
    coverage = read_csv("data_coverage_by_league_v1.csv")
    keep = ["KBO/STATIZ snapshot", "MLB Savant raw 2023", "MLB Savant raw 2024", "MLB Savant raw 2025", "MLB Savant raw 2026", "MiLB current/historical output", "NPB official stats output", "NPB+CPBL Asian quota market output"]
    rows = []
    for _, row in coverage[coverage["league_bucket"].isin(keep)].iterrows():
        rows.append([row["league_bucket"], fmt_int(row["files"]), fmt_int(row["rows"]), row["detail"]])
    add_academic_table(document, "표 2.1: 데이터 구성 요약", ["데이터", "파일", "행", "설명"], rows, widths=[4.5, 1.4, 2, 6])
    add_paragraph(
        document,
        "최종 모델 입력에서는 기사, 인터뷰, 뉴스 텍스트를 제외했다. 텍스트 자료는 문제 배경을 이해하는 보조 자료로는 사용할 수 있지만, "
        "후보 점수 산출에는 숫자형 성적 데이터와 구조화된 로스터/시장 접근성 데이터만 사용했다.",
        first_line=True,
    )

    add_heading(document, "2.2 6단계 데이터 마이닝 프레임", 2)
    frame_rows = [
        ["1", "SSG hidden weakness mining", "상황별/역할별 split으로 SSG 고유 병목 도출"],
        ["2", "KBO archetype mining", "과거 외인 성공/실패 유형과 실패 패턴 학습"],
        ["3", "Candidate market construction", "MLB/AAA/AA/NPB/CPBL 후보와 로스터 접근성 구축"],
        ["4", "KBO translation model", "해외 feature가 KBO 성공/실패로 번역될 가능성 추정"],
        ["5", "Failure risk model", "계약, 메디컬, 로스터, translation uncertainty 반영"],
        ["6", "SSG fit ranking", "팀 약점과 후보 feature를 결합해 최종 후보 산출"],
    ]
    add_academic_table(document, "표 2.2: 6단계 분석 프레임", ["단계", "이름", "역할"], frame_rows, widths=[1.3, 4.5, 8])

    add_heading(document, "2.3 통합 모델 구조", 2)
    pipeline_rows = [
        ["Input", "KBO 상황별 split, MLB Savant, MiLB, 로스터, 후보시장, 계약 규정", "숫자형/구조화 데이터만 사용"],
        ["Mining layer", "SSG 약점, KBO 외인 성공/실패 archetype, 후보 시장, KBO 번역, 실패 리스크", "각 layer가 후보를 통과/보류/탈락으로 분류"],
        ["Output", "외인타자 3명, 외인투수 진단 리드, 아시아쿼터 보류", "추천 강도를 model-supported/diagnostic/hold로 구분"],
    ]
    add_academic_table(document, "표 2.3: 통합 모델 입출력", ["구간", "내용", "출력 규칙"], pipeline_rows, widths=[2.2, 6.0, 5.2], size=8.2)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(charts["pipeline"]), width=Inches(6.2))
    add_caption(document, "그림 2.1: 통합 데이터 마이닝 파이프라인")
    add_paragraph(
        document,
        "이 구조의 핵심은 후보 점수를 바로 최종 추천으로 읽지 않는다는 점이다. 모델은 후보를 줄이는 장치이고, "
        "최종 의사결정은 contract, medical, assignment, Korea-willingness gate를 붙인 뒤 내려야 한다.",
        first_line=True,
    )

    add_heading(document, "2.4 모델 검증과 해석 경계", 2)
    decisions = read_csv("kbo_translation_failure_feature_family_decisions_v0_3.csv")
    rows = []
    for _, row in decisions.iterrows():
        rows.append([row["role_model_family"], row["target"], row["feature_family"], row["model"], fmt_num(row["mean_auc"], 3), fmt_num(row["brier_lift_vs_role_prior"], 3), row["promotion_status"]])
    add_academic_table(document, "표 2.4: 모델 검증 요약", ["Role", "Target", "Feature", "Model", "AUC", "Brier lift", "Status"], rows, widths=[1.8, 1.7, 3, 2.5, 1.2, 1.7, 2.3], size=8)
    add_paragraph(
        document,
        "검증 결과 타자 Savant classifier는 성공/실패 양쪽에서 pilot-promote를 통과했다. 반면 투수 모델은 success가 watch 수준이고 failure는 do-not-promote이므로, "
        "투수 후보는 최종 추천이 아니라 진단 리드로 해석한다.",
        first_line=True,
    )


def add_results(document: Document, charts: dict[str, Path]) -> None:
    add_heading(document, "3. 결과", 1)
    add_heading(document, "3.1 SSG 숨은 약점: 타격 runway", 2)
    runway = read_csv("ssg_2026_runway_gap_by_team.csv")
    ssg = runway[runway["t_code_name"].eq("SSG")].iloc[0]
    rows = []
    for _, row in runway.head(5).iterrows():
        rows.append([row["t_code_name"], fmt_num(row["OPS_risp"], 3), fmt_int(row["risp_ops_rank"]), fmt_num(row["OPS_on_first"], 3), fmt_int(row["on_first_ops_rank"]), fmt_num(row["OBP_on_first"], 3), fmt_num(row["risp_minus_on_first_ops"], 3)])
    add_academic_table(document, "표 3.1: SSG 상황별 runway gap", ["팀", "RISP OPS", "RISP 순위", "1루 OPS", "1루 순위", "1루 OBP", "Gap"], rows, widths=[1.7, 1.8, 1.4, 1.8, 1.4, 1.7, 1.4], size=8)
    add_paragraph(document, f"정량 해석: SSG는 RISP OPS {ssg['OPS_risp']:.3f}로 1위지만, 1루 주자 상황 OPS는 {ssg['OPS_on_first']:.3f}로 10위다. 이는 득점권 해결력보다 1루 주자를 득점권으로 옮기는 과정이 더 큰 병목임을 뜻한다.", first_line=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(charts["runway"]), width=Inches(5.4))
    add_caption(document, "그림 3.1: SSG 타격 runway gap")

    add_heading(document, "3.2 SSG 숨은 약점: 투수 import-slot inversion", 2)
    context = read_csv("ssg_pitching_message_v0_2_context_validation.csv").head(6)
    rows = []
    for _, row in context.iterrows():
        rows.append([row["context_label"], fmt_int(row["G"]), fmt_num(row["ERA"], 2), fmt_num(row["WHIP"], 2), fmt_num(row["OPS"], 3), fmt_num(row["bb9"], 2), fmt_num(row["hr9"], 2), fmt_num(row["bad_context_score"], 2)])
    add_academic_table(document, "표 3.2: SSG 투수 worst context", ["Context", "G", "ERA", "WHIP", "OPS", "BB9", "HR9", "Score"], rows, widths=[3.5, 1, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3], size=8)
    impact = read_csv("ssg_2026_import_slot_pitching_impact.csv")
    impact_rows = []
    for _, row in impact.iterrows():
        impact_rows.append([row["import_slot_group"], row["pitch_role"], fmt_int(row["starts"]), fmt_num(row["ip_per_game"], 2), fmt_num(row["era"], 2), fmt_num(row["whip"], 2), fmt_num(row["bb9"], 2), fmt_num(row["ops_allowed"], 3)])
    add_academic_table(document, "표 3.3: 외국인 투수 슬롯과 국내 투수 비교", ["Group", "Role", "Starts", "IP/G", "ERA", "WHIP", "BB9", "OPS allowed"], impact_rows, widths=[3, 1.6, 1.2, 1.3, 1.3, 1.3, 1.3, 1.8], size=8)
    add_paragraph(document, "정량 해석: import-slot starter는 ERA 6.17, WHIP 1.73, 피OPS .821로 국내 선발보다 더 큰 실점 위험을 보였다. SSG의 외인투수 보강은 탈삼진 upside보다 초반 traffic과 볼넷/피홈런을 줄이는 starter 안정화가 우선이다.", first_line=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(charts["pitch_context"]), width=Inches(5.4))
    add_caption(document, "그림 3.2: SSG 투수 bad-context score")

    add_heading(document, "3.3 KBO 외국인 성공/실패 유형", 2)
    arche = read_csv("kbo_foreign_archetype_prearrival_profile_v0_2.csv")
    rows = []
    for _, row in arche[arche["archetype_name"].isin(["hitter_failure_replacement_or_low_impact", "hitter_everyday_middle_order_anchor", "pitcher_load_bearing_rotation_anchor", "pitcher_failure_replacement_or_health_risk"])].iterrows():
        rows.append([row["role_model_family"], row["archetype_name"], fmt_int(row["rows"]), fmt_pct(row["success_rate"]), fmt_pct(row["failure_rate"]), fmt_pct(row["in_season_replaced_rate"]), fmt_num(row["median_first_kbo_war"], 2)])
    add_academic_table(document, "표 3.4: KBO 외국인 archetype 요약", ["Role", "Archetype", "Rows", "Success", "Failure", "Replaced", "WAR"], rows, widths=[1.5, 5.8, 1, 1.5, 1.5, 1.5, 1.2], size=7.5)
    add_paragraph(document, "의미: KBO 외국인 성공은 가장 화려한 단일 tool보다 시즌을 버티는 역할 안정성과 연결된다. 타자는 everyday volume, 투수는 load-bearing rotation anchor가 핵심 archetype이다.", first_line=True)

    add_heading(document, "3.4 후보 시장과 현실성 gate", 2)
    market = read_csv("candidate_market_coverage_v0_3.csv")
    rows = []
    for _, row in market[market["rows"].fillna(0).gt(0)].iterrows():
        rows.append([row["slot"], fmt_int(row["rows"]), fmt_int(row["research_lead_rows"]), fmt_int(row["market_watch_rows"]), fmt_int(row["medical_hold_rows"]), fmt_int(row["recent_release_or_dfa_rows"])])
    add_academic_table(document, "표 3.5: 후보 시장 구성", ["Slot", "Rows", "Research", "Watch", "Medical", "DFA/release"], rows, widths=[4.4, 1.4, 1.5, 1.5, 1.5, 2], size=8)
    add_paragraph(document, "실행 제안: 후보 시장은 확보됐지만 exact salary, opt-out, buyout, assignment status, medical, Korea-willingness가 아직 hard gate로 완전히 붙지 않았다. 따라서 후보명은 데이터 마이닝 리드이지 계약 확정 추천이 아니다.", first_line=True)

    add_heading(document, "3.5 최종 후보와 데이터 마이닝 통과 논리", 2)
    top = read_csv("data_mining_recommendations_top3_v1.csv")
    hrows = []
    for _, row in top[top["slot"].eq("foreign_hitter")].iterrows():
        hrows.append([fmt_int(row["rank"]), row["player_name"], row["roster_team"], fmt_int(row["age"]), str(row["is_40man"]), fmt_pct(row["dm_success_prob"]), fmt_pct(row["dm_failure_prob"]), fmt_num(row["dm_margin"], 3)])
    add_academic_table(document, "표 3.6: 최종 외인타자 후보", ["Rank", "Player", "Org", "Age", "40-man", "P(success)", "P(failure)", "Margin"], hrows, widths=[1, 3, 1.2, 1, 1.4, 1.7, 1.7, 1.5], size=8)
    prows = []
    for _, row in top[top["slot"].eq("foreign_pitcher")].iterrows():
        prows.append([fmt_int(row["rank"]), row["player_name"], row["roster_team"], fmt_int(row["age"]), str(row["is_40man"]), fmt_pct(row["dm_success_prob"]), fmt_pct(row["dm_failure_prob"]), fmt_num(row["dm_margin"], 3), row["recommendation_strength"]])
    add_academic_table(document, "표 3.7: 외인투수 진단 리드", ["Rank", "Player", "Org", "Age", "40-man", "P(success)", "P(failure)", "Margin", "Strength"], prows, widths=[0.9, 2.8, 1.2, 1, 1.4, 1.5, 1.5, 1.3, 2.4], size=7.8)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(charts["candidates"]), width=Inches(5.6))
    add_caption(document, "그림 3.3: 후보별 success/failure probability")
    add_paragraph(document, "해석: 타자는 Matos/Jones/Carlson이 모델상 낮은 실패확률과 높은 성공확률을 동시에 보였다. 투수는 Wilson/Gomber가 진단 리드이나, 모델 검증력이 watch 수준이므로 최종 추천으로 단정하지 않는다.", first_line=True)

    document.add_page_break()
    add_heading(document, "3.6 탈락 후보 audit", 2)
    hitter_pool = read_csv("data_mining_hitter_candidates_v1.csv")
    pitcher_pool = read_csv("data_mining_pitcher_candidates_v1.csv")

    def candidate_row(df: pd.DataFrame, name: str) -> pd.Series:
        return df[df["player_name"].eq(name)].iloc[0]

    decline_rows = []
    for name, reason in [
        ("Drew Waters", "top3보다 margin이 낮고 실패확률이 높아 예비군으로 하향"),
        ("Dylan Moore", "나이와 K%, margin 하락으로 first-base converter 우선순위에서 밀림"),
        ("Jack Suwinski", "barrel은 있으나 failure probability가 높아 SSG fit에서 탈락"),
        ("Luis Arráez", "모델 점수는 높지만 40-man/MLB active low-access gate에서 탈락"),
    ]:
        row = candidate_row(hitter_pool, name)
        decline_rows.append([name, "타자", fmt_pct(row["dm_success_prob"]), fmt_pct(row["dm_failure_prob"]), fmt_num(row["dm_margin"], 3), reason])
    for name, reason in [
        ("Dietrich Enns", "sample gate는 통과했으나 margin 음수와 비용 block으로 보류"),
        ("Kris Bubic", "수치상 upside가 있어도 injured/40-man gate 때문에 release 불가"),
    ]:
        row = candidate_row(pitcher_pool, name)
        decline_rows.append([name, "투수", fmt_pct(row["dm_success_prob"]), fmt_pct(row["dm_failure_prob"]), fmt_num(row["dm_margin"], 3), reason])
    add_academic_table(
        document,
        "표 3.8: 주요 탈락 후보 audit",
        ["선수", "Slot", "P(success)", "P(failure)", "Margin", "탈락/보류 사유"],
        decline_rows,
        widths=[2.6, 1.2, 1.6, 1.6, 1.4, 5.8],
        size=7.5,
    )
    add_paragraph(
        document,
        "이 audit의 의미는 모델이 단순히 점수가 높은 선수를 고르는 것이 아니라 접근성, 실패확률, SSG fit, 비용 block을 "
        "동시에 적용해 후보를 줄였다는 점이다. 특히 Arráez처럼 모델 score가 좋아도 실무 gate가 막히면 최종 후보가 될 수 없다.",
        first_line=True,
    )

    document.add_page_break()
    add_heading(document, "3.7 후보별 통과 조건과 반증 조건", 2)
    decision_rows = [
        ["Luis Matos", "타자 lead", "non-40man, PA 표본, 낮은 실패확률, OF fit", "현재 조직 통제, 이적 의사, KBO행 가능성", "우선 접촉 전 contract/assignment 확인"],
        ["Nolan Jones", "타자 lead", "높은 성공확률, 좌타 OF/DH runway 보강 가능성", "공개 비용 신호와 접근성", "비용 gate 통과 시 1차 후보 유지"],
        ["Dylan Carlson", "타자 reserve", "모델상 3순위, 포지션 유연성", "현재 조직/계약 소스 불일치", "소스 재확인 전 확정 표현 금지"],
        ["Bryse Wilson", "투수 diagnostic", "MiLB IP, BB9/HR9, non-40man gate", "선발 지속성, 최근 구위, 80-90구 가능성", "영상/스플릿 검증 후 보드 유지"],
        ["Austin Gomber", "투수 diagnostic", "workload continuity와 선발 경험", "HR9, 비용, opt-out", "장타 억제 근거가 없으면 하향"],
        ["Dietrich Enns", "hold", "표본과 K9는 존재", "margin 음수, 비용 block", "추천이 아니라 보류/반례로 처리"],
    ]
    add_academic_table(
        document,
        "표 3.9: 후보별 통과 및 반증 조건",
        ["선수", "판정", "통과 근거", "반증 조건", "다음 액션"],
        decision_rows,
        widths=[2.5, 2.0, 3.4, 3.3, 3.2],
        size=7.4,
    )
    add_paragraph(
        document,
        "실무 해석: 이 표의 목적은 후보를 홍보하는 것이 아니라 후보를 탈락시킬 조건을 명확히 하는 것이다. "
        "특히 투수 후보는 모델 검증력이 낮으므로, 반증 조건을 통과하지 못하면 최종 보드에서 제외해야 한다.",
        first_line=True,
    )

    add_heading(document, "3.8 반박 가능성 점검", 2)
    objection_rows = [
        ["SSG는 그냥 장타 외야수가 필요한 것 아닌가?", "RISP OPS 1위, 1루 주자 OPS 10위", "장타 총량보다 first-base traffic 전환력이 우선"],
        ["득점권 성적은 운 아닌가?", "문제 지점은 득점권이 아니라 득점권 이전 단계", "운 논쟁을 피하고 주자 상황별 process로 해석"],
        ["투수는 삼진을 많이 잡는 선수가 최고 아닌가?", "SSG worst context는 early, runner-on-base, RISP, BB9/HR9", "탈삼진보다 traffic command와 starter length가 우선"],
        ["후보가 너무 MLB급이라 현실성이 낮지 않나?", "non-40man gate는 통과했지만 salary/assignment는 미완성", "계약 gate 전에는 영입 확정 추천이 아니라 데이터 리드"],
        ["왜 아시아쿼터 후보는 빠졌나?", "국적, 최근 아시아리그 소속, 200k cap을 선수별로 붙여야 함", "아시아쿼터는 별도 hard-gate shortlist가 필요"],
    ]
    add_academic_table(
        document,
        "표 3.10: 주요 반박과 대응 근거",
        ["반박", "데이터 근거", "보고서 대응"],
        objection_rows,
        widths=[4.2, 4.5, 5.0],
        size=7.6,
    )


def add_discussion_limit_conclusion(document: Document) -> None:
    add_heading(document, "4. 논의", 1)
    add_heading(document, "4.1 실무 적용 원칙", 2)
    add_paragraph(document, "타자 보강은 '외야 장타'라는 일반론보다 우투수 상대 1루 주자 상황에서 죽지 않고 다음 베이스를 만드는 OF/DH 유형으로 좁혀야 한다. 투수 보강은 구속/탈삼진보다 5이닝 초반을 6이닝으로 바꾸는 traffic-command starter에 초점을 둬야 한다.", first_line=True)
    add_heading(document, "4.2 후보별 운영", 2)
    rows = [
        ["Luis Matos", "타자 1순위", "모델 margin 최상. 저비용 가능성 및 조직 통제 확인 필요"],
        ["Nolan Jones", "타자 2순위", "모델 강도 높음. 비용/접근성 gate가 핵심 반증 포인트"],
        ["Dylan Carlson", "타자 3순위", "예비 3순위. 현재 조직/계약 소스 재확인 필요"],
        ["Bryse Wilson", "투수 진단 1순위", "선발 지속성과 command/damage control 영상 검증 필요"],
        ["Austin Gomber", "투수 진단 2순위", "workload 장점. HR9와 비용/옵트아웃 확인 필요"],
        ["Dietrich Enns", "보류", "margin 음수와 비용 block 성격으로 추천 아님"],
    ]
    add_academic_table(document, "표 4.1: 후보별 운영 판단", ["선수", "판정", "실무 해석"], rows, widths=[3, 2.5, 8], size=8)

    add_heading(document, "4.3 공개 연봉 및 계약 현실성 gate", 2)
    salary_rows = [
        ["Luis Matos", "외인타자", "$129k-$780k range", "저비용 가능성은 있으나 조직 통제/이적 의사 확인 필요"],
        ["Nolan Jones", "외인타자", "$2.0M signal", "비용과 접근성 gate가 가장 큰 반증 포인트"],
        ["Dylan Carlson", "외인타자", "$2.0M signal / minor terms unclear", "현재 조직과 보장액 재확인 전 확정 금지"],
        ["Bryse Wilson", "외인투수", "$850k signal", "진단 리드이나 salary/role 확인 필요"],
        ["Austin Gomber", "외인투수", "minor deal undisclosed", "접근성은 있으나 정확 비용 미확인"],
        ["Dietrich Enns", "외인투수", "$2.5M-$2.625M signal", "비용 block과 negative margin이 동시에 존재"],
    ]
    add_academic_table(
        document,
        "표 4.2: 공개 연봉 및 계약 현실성 gate",
        ["선수", "Slot", "공개 비용 신호", "실무 해석"],
        salary_rows,
        widths=[2.7, 1.9, 3.4, 6.0],
        size=7.6,
    )
    add_paragraph(
        document,
        "해석 원칙: 공개 연봉은 아직 전 후보군 전체에 완전하게 붙지 않았으므로 모델 핵심 input이 아니라 후속 hard gate로 둔다. "
        "따라서 모델 점수가 높아도 보장계약, opt-out, buyout, assignment status가 불리하면 최종 후보에서 제외해야 한다.",
        first_line=True,
    )

    add_heading(document, "4.4 아시아쿼터 잠금 사유", 2)
    asian_rows = [
        ["국적", "BFA 국가 또는 호주 국적 필요", "선수별 nationality/passport 확인 필요"],
        ["리그 이력", "전년도 또는 당해년도 아시아리그 소속 필요", "NPB/CPBL/KBO 이력 매칭 필요"],
        ["비용", "신규 영입 총액 200k USD cap", "연봉, 옵션, 이적료, buyout을 같이 계산"],
        ["인원", "최대 1명", "외인투수/타자 보강과 roster 전략 동시 검토"],
    ]
    add_academic_table(
        document,
        "표 4.3: 아시아쿼터 hard gate",
        ["Gate", "규정/조건", "데이터 요구"],
        asian_rows,
        widths=[2.2, 5.5, 5.8],
        size=8,
    )
    add_paragraph(
        document,
        "따라서 아시아쿼터는 현재 보고서에서 이름을 확정하지 않는 것이 맞다. 후보 pool은 존재하지만 국적, 리그 이력, 200k cap을 "
        "후보별로 통과시킨 뒤 별도 shortlist로 제시해야 한다.",
        first_line=True,
    )

    document.add_page_break()
    add_heading(document, "4.5 민감도 및 강건성 점검", 2)
    hitter_pool = read_csv("data_mining_hitter_candidates_v1.csv")
    pitcher_pool = read_csv("data_mining_pitcher_candidates_v1.csv")
    hitter_gate = hitter_pool[hitter_pool["data_mining_gate_pass"].eq(True)].copy()
    pitcher_gate = pitcher_pool[pitcher_pool["data_mining_gate_pass"].eq(True)].copy()

    def top_names(df: pd.DataFrame, score: pd.Series, n: int) -> str:
        tmp = df.assign(_score=score).sort_values("_score", ascending=False).head(n)
        return ", ".join(tmp["player_name"].tolist())

    h_base = hitter_gate["dm_margin"]
    h_failure = hitter_gate["dm_success_prob"] - 1.2 * hitter_gate["dm_failure_prob"]
    h_success = 0.85 * hitter_gate["dm_success_prob"] - hitter_gate["dm_failure_prob"]
    h_sample = hitter_gate["dm_margin"] - ((250 - hitter_gate["recent_pa"].clip(upper=250)) / 250) * 0.15

    p_base = pitcher_gate["dm_margin"]
    p_failure = pitcher_gate["dm_success_prob"] - 1.2 * pitcher_gate["dm_failure_prob"]
    p_success = 0.85 * pitcher_gate["dm_success_prob"] - pitcher_gate["dm_failure_prob"]
    p_command = (
        pitcher_gate["dm_margin"]
        - pitcher_gate["pre_kbo_milb_bb9"].gt(3).astype(float) * 0.03
        - pitcher_gate["pre_kbo_milb_hr9"].gt(1.2).astype(float) * 0.03
        + pitcher_gate["pre_kbo_milb_ip"].gt(100).astype(float) * 0.02
    )
    sensitivity_rows = [
        ["Base margin", top_names(hitter_gate, h_base, 3), top_names(pitcher_gate, p_base, 2), "기본 결론"],
        ["Failure +20%", top_names(hitter_gate, h_failure, 3), top_names(pitcher_gate, p_failure, 2), "타자 top3 유지"],
        ["Success -15%", top_names(hitter_gate, h_success, 3), top_names(pitcher_gate, p_success, 2), "타자 top3 유지"],
        ["Sample/command stress", top_names(hitter_gate, h_sample, 3), top_names(pitcher_gate, p_command, 2), "투수는 lead만 유지"],
    ]
    add_academic_table(
        document,
        "표 4.4: 민감도 및 강건성 점검",
        ["Stress scenario", "외인타자 top3", "외인투수 lead", "해석"],
        sensitivity_rows,
        widths=[3.0, 4.6, 3.3, 3.0],
        size=7.5,
    )
    add_paragraph(
        document,
        "민감도 해석: 타자 보드는 실패확률 가중, 성공확률 할인, 표본 penalty를 적용해도 top3가 유지된다. "
        "반면 투수 보드는 Wilson/Gomber 순서는 유지되지만 margin 자체가 작으므로 확정 추천이 아니라 진단 리드로 표현해야 한다.",
        first_line=True,
    )

    add_heading(document, "4.6 최종 의사결정 gate", 2)
    gate_rows = [
        ["1. Model pass", "hitter classifier 또는 pitcher diagnostic 통과", "타자 3명 통과, 투수 2명 진단 통과", "데이터 후보 유지"],
        ["2. Market access", "non-40man, DFA/release/minor status, 조직 통제 확인", "현재 1차 gate만 통과", "assignment 재확인 전 확정 금지"],
        ["3. Contract economics", "salary, guarantee, opt-out, buyout, 이적료 확인", "부분 공개 비용만 반영", "계약표 완성 전 순위 고정 금지"],
        ["4. Medical/role", "부상, 80-100구 가능성, 수비 포지션, workload 확인", "정량 데이터 외 수동 검증 필요", "스카우팅 카드 필수"],
        ["5. Korea-willingness", "선수/에이전트의 KBO행 의사", "현재 미확인", "최종 접촉 전 hard gate"],
        ["6. Final board", "위 5개 gate 통과 후 3+3 보드 확정", "현재는 research board", "발표에서는 '최종 영입 확정' 표현 금지"],
    ]
    add_academic_table(
        document,
        "표 4.5: 최종 의사결정 gate",
        ["Gate", "통과 기준", "현재 상태", "결정 규칙"],
        gate_rows,
        widths=[3.0, 4.3, 3.5, 3.0],
        size=7.5,
    )

    document.add_page_break()
    add_heading(document, "5. 한계", 1)
    rows = [
        ["데이터 시점", "STATIZ 최신 스냅샷이 2026-06-11 기준", "최신 경기 반영 후 Layer 1 재검증"],
        ["투수 모델", "pitcher classifier가 watch 수준", "과거 KBO 외인투수 pre-arrival 데이터 추가 backfill"],
        ["연봉/계약", "전 후보군 exact salary/opt-out 미완성", "Top 20에 current salary, guarantee, buyout 부착"],
        ["아시아쿼터", "국적/리그 이력/200k cap gate 미완성", "별도 shortlist 생성"],
        ["수동 스카우팅", "영상/메디컬/선수 의향은 미반영", "최종 영입 전 human review 필수"],
    ]
    add_academic_table(document, "표 5.1: 한계와 후속 작업", ["영역", "한계", "후속 작업"], rows, widths=[2.5, 5.5, 5.5], size=8)

    add_heading(document, "6. 결론", 1)
    add_paragraph(document, "본 보고서의 결론은 SSG가 외국인 선수 보강을 포지션/툴 단위가 아니라 game-state 단위로 재정의해야 한다는 것이다. 타자는 first-base traffic converter, 투수는 load-bearing traffic-command starter가 핵심이며, 이를 KBO 외인 성공/실패 데이터에 통과시킨 결과 타자는 Matos/Jones/Carlson이 우선 보드에 오른다. 투수는 Wilson/Gomber를 진단 리드로 두되, 현재 모델 수준에서는 확정 추천이 아니라 추가 검증 우선순위로 제시한다. 최종 영입 판단은 이 결론에 salary, opt-out, buyout, assignment, medical, Korea-willingness gate를 붙인 뒤 내려야 한다.", first_line=True)


def add_appendix(document: Document) -> None:
    add_heading(document, "부록 A. 재현 산출물", 1)
    rows = [
        ["A1", "data_mining_recommendations_top3_v1.csv", "최종 후보 3+3"],
        ["A2", "data_mining_hitter_candidates_v1.csv", "외인타자 전체 랭킹"],
        ["A3", "data_mining_pitcher_candidates_v1.csv", "외인투수 전체 랭킹"],
        ["A4", "ssg_2026_runway_gap_by_team.csv", "SSG 타격 runway gap"],
        ["A5", "ssg_pitching_message_v0_2_context_validation.csv", "SSG 투수 context 검증"],
        ["A6", "kbo_translation_failure_feature_family_decisions_v0_3.csv", "모델 검증표"],
        ["A7", "candidate_market_coverage_v0_3.csv", "후보 시장 구성"],
        ["A8", "kbo_contract_constraints_v1.csv", "KBO 계약/규정 gate"],
    ]
    add_academic_table(document, "표 A.1: 주요 산출물", ["ID", "파일", "역할"], rows, widths=[1, 6, 6], size=8)


def build_doc() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    charts = make_charts()
    document = Document()
    setup_document(document)
    set_header_footer(document.sections[0], header_on=False, page_number=False)
    add_title_page(document)

    front = document.add_section(WD_SECTION.NEW_PAGE)
    setup_document(document)
    set_page_numbering(front, start=1, fmt="lowerRoman")
    set_header_footer(front, header_on=False)
    add_abstract(document)
    add_executive_summary(document)
    document.add_page_break()
    add_contents(document)

    body = document.add_section(WD_SECTION.NEW_PAGE)
    setup_document(document)
    set_page_numbering(body, start=1, fmt="decimal")
    set_header_footer(body, header_on=True)
    add_intro(document)
    add_methods(document, charts)
    add_results(document, charts)
    add_discussion_limit_conclusion(document)
    add_appendix(document)

    document.save(DOC_PATH)
    return DOC_PATH


if __name__ == "__main__":
    print(build_doc())
