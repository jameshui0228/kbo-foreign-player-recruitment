#!/usr/bin/env python3
"""Create a team-shareable DOCX summary for the SSG recruitment project."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "tables"
DOC_OUT_DIR = ROOT / "output" / "doc"
DOC_PATH = DOC_OUT_DIR / "SSG_외국인선수_데이터마이닝_후보추천_v1.docx"

TOP_PATH = OUT / "data_mining_recommendations_top3_v1.csv"
AUDIT_PATH = OUT / "data_mining_model_audit_v1.csv"
DECISIONS_PATH = OUT / "kbo_translation_failure_feature_family_decisions_v0_3.csv"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: str = "111827") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True, size=8, color="FFFFFF")
        set_cell_shading(header_cells[i], "1F2937")
        if widths:
            header_cells[i].width = Cm(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)
            if widths:
                cells[i].width = Cm(widths[i])

    document.add_paragraph()


def add_callout(document: Document, title: str, body: str, fill: str = "EAF2FF") -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    p.add_run("\n")
    run = p.add_run(body)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    document.add_paragraph()


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        if level == 1:
            run.font.color.rgb = RGBColor(17, 24, 39)
        else:
            run.font.color.rgb = RGBColor(31, 41, 55)


def add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        rest = text[len(bold_prefix):]
        run = p.add_run(rest)
    else:
        run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")


def pct(value: float) -> str:
    return f"{float(value) * 100:.1f}%"


def build_doc() -> Path:
    top = pd.read_csv(TOP_PATH)
    audit = pd.read_csv(AUDIT_PATH)
    decisions = pd.read_csv(DECISIONS_PATH)

    DOC_OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SSG 외국인 선수 영입 후보 추천\n데이터 마이닝 결론 v1")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SDA 2기 SSG 랜더스 대체 외국인/아시아쿼터 선수 영입 프로젝트")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(75, 85, 99)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")

    add_callout(
        document,
        "한 줄 결론",
        "타자는 데이터 마이닝 모델이 Luis Matos, Nolan Jones, Dylan Carlson을 강하게 밀어 올렸다. "
        "투수는 현재 데이터만으로 확정 추천까지는 약하며, Bryse Wilson과 Austin Gomber를 진단 리드로 두고 연봉/계약 데이터를 추가해야 한다.",
        fill="DDEBFF",
    )

    add_heading(document, "1. 분석 원칙", 1)
    add_bullet(document, "기사, 인터뷰, 뉴스, 텍스트 기반 변수는 최종 모델 입력에서 제외했다.")
    add_bullet(document, "과거 KBO 외국인 선수 성공/실패 라벨과 입단 전 숫자형 지표만 학습 데이터로 사용했다.")
    add_bullet(document, "현재 후보군에는 MLB/MiLB/Savant 숫자형 지표, 40인 로스터 여부, 시장 접근성 구조화 데이터를 적용했다.")
    add_bullet(document, "최종 후보는 좋은 선수 순위가 아니라, 과거 KBO 성공/실패 패턴을 통과한 후보로 해석한다.")

    add_heading(document, "2. 핵심 모델 구조", 1)
    add_paragraph(
        document,
        "이번 버전은 기존의 수동 가중치 점수표가 아니라, 과거 KBO 외국인 선수 데이터를 학습한 분류 모델을 현재 후보 시장에 적용한 방식이다.",
    )
    add_table(
        document,
        ["구분", "학습 행", "Feature family", "모델", "Target", "판정"],
        [
            ["외인타자", "22", "Savant", "ridge logit", "success", "promoted"],
            ["외인타자", "22", "Savant", "ridge logit", "failure", "promoted"],
            ["외인투수", "49", "MiLB damage/command", "sparse L1", "success", "watch"],
            ["외인투수", "49", "MiLB damage/command", "sparse L1", "failure", "warning only"],
        ],
        widths=[2.2, 1.5, 3.2, 2.5, 2.0, 3.2],
    )

    add_heading(document, "3. 검증 결과", 1)
    promoted = decisions[decisions["promotion_status"].isin(["pilot_promote", "watch"])].copy()
    validation_rows = []
    for row in promoted.to_dict("records"):
        validation_rows.append(
            [
                "타자" if row["role_model_family"] == "hitter" else "투수",
                row["target"],
                "Savant" if row["feature_family"] == "savant_only" else "MiLB dmg/cmd",
                row["model"],
                f"{row['mean_auc']:.3f}",
                f"{row['brier_lift_vs_role_prior']:.3f}",
                row["promotion_status"],
            ]
        )
    add_table(
        document,
        ["Role", "Target", "Feature", "Model", "AUC", "Brier lift", "Status"],
        validation_rows,
        widths=[1.5, 1.6, 3.0, 2.4, 1.4, 1.8, 2.2],
    )
    add_callout(
        document,
        "해석",
        "타자 모델은 성공/실패 양쪽에서 pilot_promote를 통과했다. 따라서 외인타자 후보 추천은 모델 근거가 비교적 강하다. "
        "반면 투수 모델은 success만 watch 수준이고 failure는 promoted가 아니므로, 투수 후보는 확정 추천이 아니라 추가 검증 리스트로 봐야 한다.",
        fill="FFF7D6",
    )

    add_heading(document, "4. 외국인 타자 최종 후보", 1)
    hitter = top[top["slot"].eq("foreign_hitter")].copy()
    hitter_rows = []
    for row in hitter.to_dict("records"):
        hitter_rows.append(
            [
                int(row["rank"]),
                row["player_name"],
                row["roster_team"],
                int(row["age"]),
                str(row["is_40man"]),
                pct(row["dm_success_prob"]),
                pct(row["dm_failure_prob"]),
                f"{row['dm_margin']:.3f}",
            ]
        )
    add_table(
        document,
        ["순위", "선수", "조직", "나이", "40인", "성공확률", "실패확률", "Margin"],
        hitter_rows,
        widths=[1.2, 3.0, 1.4, 1.2, 1.3, 1.8, 1.8, 1.7],
    )
    add_paragraph(document, "타자 결론: ", bold_prefix="타자 결론: ")
    add_bullet(document, "Luis Matos는 성공확률 92.4%, 실패확률 8.2%로 가장 안정적인 모델 후보로 나왔다.")
    add_bullet(document, "Nolan Jones는 성공확률 90.2%로 강한 후보지만, 실제 연봉/계약 가능성 확인이 필요하다.")
    add_bullet(document, "Dylan Carlson은 1, 2순위보다는 약하지만 여전히 모델상 유의미한 외야 후보로 남는다.")
    add_bullet(document, "이전 점수표에서 높았던 Will Brennan, Dominic Fletcher는 데이터 마이닝 모델에서 표본/성공패턴 신뢰도가 부족해 내려갔다.")

    add_heading(document, "5. 외국인 투수 진단 후보", 1)
    pitcher = top[top["slot"].eq("foreign_pitcher")].copy()
    pitcher_rows = []
    for row in pitcher.to_dict("records"):
        pitcher_rows.append(
            [
                int(row["rank"]),
                row["player_name"],
                row["roster_team"],
                int(row["age"]),
                str(row["is_40man"]),
                pct(row["dm_success_prob"]),
                pct(row["dm_failure_prob"]),
                f"{row['dm_margin']:.3f}",
                "진단 리드" if row["recommendation_strength"] == "diagnostic_lead" else "보류",
            ]
        )
    add_table(
        document,
        ["순위", "선수", "조직", "나이", "40인", "성공확률", "실패위험", "Margin", "강도"],
        pitcher_rows,
        widths=[1.0, 2.7, 1.2, 1.1, 1.1, 1.7, 1.7, 1.5, 3.1],
    )
    add_paragraph(document, "투수 결론: ", bold_prefix="투수 결론: ")
    add_bullet(document, "Bryse Wilson과 Austin Gomber는 시장/표본 게이트를 통과한 진단 리드다.")
    add_bullet(document, "Dietrich Enns는 게이트는 통과했지만 모델 margin이 음수이므로 추천이 아니라 보류로 봐야 한다.")
    add_bullet(document, "투수 모델은 현재 AUC 0.603 수준이라 타자 모델만큼 강하지 않다. 추가 데이터 없이 확정 추천하면 과장이다.")

    add_heading(document, "6. 팀 의사결정용 정리", 1)
    add_callout(
        document,
        "현재 확정 가능한 주장",
        "외국인 타자는 Luis Matos, Nolan Jones, Dylan Carlson 중심으로 검토하는 것이 가장 방어 가능하다. "
        "투수는 Bryse Wilson, Austin Gomber를 우선 진단하되, 실제 연봉/보장계약/최근 건강 상태가 붙기 전에는 최종 후보로 못 박지 않는다.",
        fill="E8F5E9",
    )
    add_heading(document, "7. 다음 액션", 1)
    add_bullet(document, "Spotrac, Baseball Prospectus, MLB payroll/transaction 데이터를 활용해 실제 연봉 및 보장계약 정보를 후보군에 붙인다.")
    add_bullet(document, "Nolan Jones, Dylan Carlson처럼 모델상 강하지만 실제 계약 접근성이 애매할 수 있는 후보를 비용 필터로 재검증한다.")
    add_bullet(document, "투수 쪽은 과거 KBO 외인투수 pre-arrival 데이터와 현재 후보 MiLB 표본을 더 확장한 뒤 재학습한다.")
    add_bullet(document, "최종 모델은 성공확률, 실패확률, 40인 여부, 실제 연봉, SSG fit을 모두 통과한 후보만 남기는 방식으로 정리한다.")

    add_heading(document, "8. 산출물 위치", 1)
    add_bullet(document, "최종 후보 CSV: outputs/tables/data_mining_recommendations_top3_v1.csv")
    add_bullet(document, "전체 타자 후보 랭킹: outputs/tables/data_mining_hitter_candidates_v1.csv")
    add_bullet(document, "전체 투수 후보 랭킹: outputs/tables/data_mining_pitcher_candidates_v1.csv")
    add_bullet(document, "모델 재현 스크립트: scripts/build_data_mining_recommendation_model_v1.py")

    section = document.add_section(WD_SECTION.CONTINUOUS)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    document.save(DOC_PATH)
    return DOC_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
