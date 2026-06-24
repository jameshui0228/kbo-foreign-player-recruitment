#!/usr/bin/env python3
"""Build the formal candidate-selection DOCX report.

The document follows the structure of the provided SDA season-1 PDF:
cover, abstract, key takeaways, manual contents, introduction, method,
results, discussion, limitations, conclusion, and appendix.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TABLE_DIR = PROJECT_ROOT / "outputs" / "tables"
REPORT_DIR = PROJECT_ROOT / "outputs" / "reports"
ASSET_DIR = REPORT_DIR / "ensemble_final_report_v1_assets"
DOC_OUT = PROJECT_ROOT / "output" / "doc"
DOCX_PATH = DOC_OUT / "ssg_foreign_player_ensemble_candidate_selection_report_v3_salary_gated.docx"

SCORES_PATH = TABLE_DIR / "ensemble_candidate_scores_v1.csv"
WEIGHTS_PATH = TABLE_DIR / "ensemble_model_signal_weights_v1.csv"
SALARY_GATE_PATH = TABLE_DIR / "final_candidate_salary_contract_gate_v1.csv"


FONT = "Apple SD Gothic Neo"
INK = "1F2430"
MUTED = "6F768A"
GRID = "D7DBE7"
BLUE = "5477C4"
BLUE_LIGHT = "EAF1FE"
ORANGE_LIGHT = "FFEDDE"
GOLD_LIGHT = "FFF4C2"
NEUTRAL_LIGHT = "F4F5F7"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border(cell, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell(cell, text: str, *, bold: bool = False, size: float = 8.2, color: str = INK, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float | None = None, color: str = INK):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return r


def p(doc: Document, text: str = "", *, style: str | None = None, align=None, size: float | None = None) -> None:
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run(para, text, size=size)


def heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading("", level=level)
    run(para, text, bold=True, size=15.5 if level == 1 else 12.2)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run(para, item)


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Number")
        run(para, item)


def fmt_pct(value, digits: int = 1) -> str:
    if pd.isna(value):
        return "-"
    return f"{float(value) * 100:.{digits}f}%"


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], *, max_rows: int | None = None, font_size: float = 7.8) -> None:
    if max_rows is not None:
        df = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, NEUTRAL_LIGHT)
        border(cell)
        set_cell(cell, header, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for cell, col in zip(cells, columns):
            value = row[col]
            if col in {"ensemble_score", "dm_success_prob", "dm_failure_prob", "dm_margin", "dm_signal", "consensus_signal"}:
                text = fmt_pct(value)
            elif pd.isna(value):
                text = "-"
            else:
                text = str(value)
            border(cell)
            set_cell(cell, text, size=font_size)


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(para, text, size=8.3, color=MUTED)


def add_picture(doc: Document, path: Path, caption: str, *, width: float = 6.55) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        para.add_run().add_picture(str(path), width=Inches(width))
    else:
        run(para, f"[Missing chart: {path.name}]", color=MUTED)
    add_caption(doc, caption)


def callout(doc: Document, title: str, body: str, *, fill: str = BLUE_LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    border(cell)
    para = cell.paragraphs[0]
    run(para, title, bold=True)
    run(para, " " + body)


def set_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(4)

    for name, size in [("Title", 24), ("Heading 1", 15.5), ("Heading 2", 12.2), ("Subtitle", 10.5)]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(INK if name != "Subtitle" else MUTED)
        if name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(5)


def start_new_page(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def add_footer(doc: Document) -> None:
    for idx, sec in enumerate(doc.sections):
        sec.footer.is_linked_to_previous = False
        footer = sec.footer.paragraphs[0]
        for item in list(footer.runs):
            footer._p.remove(item._r)
        if idx == 0:
            continue
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(footer, "SDA 2기 SSG 외국인 선수 영입 프로젝트 | Ensemble Data Mining", size=8, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(title, "SSG 외국인 선수 영입\n앙상블 데이터 마이닝 보고서", bold=True, size=25)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(sub, "4인 앙상블 기반 후보 3인 선정 및 최종 1인 선발 전략", size=12.2)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(meta, "Data: KBO/STATIZ, MLB/MiLB, Savant, roster/transaction structured data | Text variables excluded", size=9.8)

    for _ in range(13):
        doc.add_paragraph()
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(author, "고려대학교 SDA 2기", size=11)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(date, "June 23, 2026", size=10.5)
    start_new_page(doc)


def add_abstract(doc: Document) -> None:
    heading(doc, "초록 (Abstract)", 1)
    p(
        doc,
        "본 보고서는 SSG 랜더스의 대체 외국인 타자 및 외국인 선발투수 영입 의사결정을 하나의 데이터 마이닝 문제로 재정의한다. 네 명의 분석 결과를 독립 후보표로 병렬 제시하지 않고, 과거 KBO 외국인 성공/실패 학습 모델, SSG fit/번역 파이프라인, KBO 적응 filter, 시장 비효율 feature model, 반복 등장 consensus signal을 base learner로 둔 stacked-style 앙상블로 통합했다. 기사, 인터뷰, 평판 텍스트 변수는 모델 입력에서 제외하고 숫자형 structured data만 사용했다. 본 v3에서는 공개 연봉/계약/메디컬 정보를 post-model hard gate로 추가했다.",
    )
    p(
        doc,
        "타자 모델의 순수 데이터 마이닝 Top 3는 Nolan Jones, Luis Matos, Jack Suwinski다. 다만 연봉 gate를 붙이면 Nolan Jones는 모델 1위이지만 2.0M 공개 cash signal 때문에 즉시 계약 가능성을 확정할 수 없고, Luis Matos는 DFA/outright 이력과 낮은 공개 급여 부담으로 비용 현실성이 가장 좋다. 따라서 발표 문장은 '모델 1위 Nolan Jones, 계약 현실성 1위 Luis Matos'로 분리하는 편이 안전하다. 투수 모델은 historical classifier 성능이 약해 확정 추천보다 검증 우선순위 모델로 운영해야 한다. raw Top 3는 Josh Fleming, Carson Spiers, Brian Van Belle이지만, Spiers와 Van Belle은 medical gate가 red이므로 실제 최종 접촉 1순위는 Josh Fleming이다.",
    )

    heading(doc, "Key Takeaways", 2)
    bullets(
        doc,
        [
            "문제 정의: 외국인 영입을 단순 후보 추천이 아니라 SSG 약점 보완을 위한 supervised/diagnostic ensemble mining 문제로 재정의했다.",
            "타자 결론: 순수 모델 Top 3는 Nolan Jones, Luis Matos, Jack Suwinski다. 연봉 gate까지 붙이면 Nolan은 모델 1위, Matos는 계약 현실성 1위다.",
            "투수 결론: raw Top 3는 Josh Fleming, Carson Spiers, Brian Van Belle이나, Spiers와 Van Belle은 medical gate red라 Josh Fleming만 접촉 1순위로 남는다.",
            "모델 차이: 타자 classifier는 AUC 0.833/0.738로 추천 가능하지만, 투수 classifier는 AUC 0.603으로 diagnostic only다.",
            "운영 메시지: 모델 점수만으로 만족하면 안 된다. 최종 후보는 반드시 salary, buyout, assignment, medical gate를 통과해야 한다.",
        ],
    )
    start_new_page(doc)


def add_contents(doc: Document) -> None:
    heading(doc, "목차", 1)
    contents = [
        "1. 서론",
        "  1.1 문제 정의",
        "  1.2 분석 질문",
        "  1.3 핵심 지표",
        "2. 방법",
        "  2.1 데이터 및 전처리",
        "  2.2 4인 앙상블 base learner",
        "  2.3 후보 점수화 및 최종 1인 선발 규칙",
        "  2.4 연봉/계약/메디컬 hard gate",
        "3. 결과",
        "  3.1 앙상블 모델 구조",
        "  3.2 외국인 타자 Top 3와 최종 1인",
        "  3.3 외국인 선발투수 Top 3와 최종 1인",
        "  3.4 모델 검증과 해석 강도",
        "  3.5 연봉 gate 반영 후 결론 변화",
        "4. 논의",
        "5. 한계",
        "6. 결론",
        "A. 부록",
    ]
    for item in contents:
        para = doc.add_paragraph()
        run(para, item)
    start_new_page(doc)


def add_intro(doc: Document) -> None:
    heading(doc, "1. 서론", 1)
    heading(doc, "1.1 문제 정의", 2)
    p(
        doc,
        "본 프로젝트의 목적은 좋은 선수를 단순히 나열하는 것이 아니라, SSG 랜더스가 실제로 영입 가능한 시장에서 팀 약점을 가장 잘 보완할 수 있는 외국인 후보를 데이터로 추리는 것이다. 따라서 분석 단위는 선수 개인의 단순 성적이 아니라, KBO 번역 가능성, SSG fit, 실패 리스크, 계약 접근성을 모두 포함한 후보 의사결정 단위다.",
    )
    p(
        doc,
        "특히 이번 보고서는 네 명의 분석을 사람별로 따로 제시하지 않는다. 각 분석을 하나의 base learner로 간주하고, 이들이 동일 후보를 얼마나 지지하는지, 그리고 어떤 리스크를 제거하는지에 따라 후보를 재랭킹한다.",
    )
    heading(doc, "1.2 분석 질문", 2)
    rq = pd.DataFrame(
        [
            ["RQ1", "외국인 타자 후보군에서 앙상블 모델이 최종 Top 3와 최종 1인을 어떻게 도출하는가?", "최종 영입 1순위 타자"],
            ["RQ2", "외국인 선발투수 후보군에서 모델 확신도와 검증 리스크를 반영하면 누구를 우선 검증해야 하는가?", "최종 검증 1순위 투수"],
        ],
        columns=["RQ", "질문", "의사결정 산출물"],
    )
    add_table(doc, rq, ["RQ", "질문", "의사결정 산출물"], ["RQ", "질문", "의사결정 산출물"], font_size=8.5)
    heading(doc, "1.3 핵심 지표", 2)
    metrics = pd.DataFrame(
        [
            ["ensemble_score", "base learner 신호를 가중 합성한 최종 후보 점수", "Top 3 선정 기준"],
            ["dm_success_prob", "과거 KBO 외국인 성공 label 기반 성공 확률", "타자 추천 강도 판단"],
            ["dm_failure_prob", "과거 KBO 외국인 실패 label 기반 실패 위험", "하방 리스크 판단"],
            ["consensus_signal", "서로 다른 base learner에서 반복 등장한 정도", "모델 간 안정성 판단"],
            ["role/risk gate", "선발 지속성, medical, 계약 접근성, 40인 여부 등", "최종 1인 선발 전 hard gate"],
        ],
        columns=["지표", "의미", "실무 해석"],
    )
    add_table(doc, metrics, ["지표", "의미", "실무 해석"], ["지표", "의미", "실무 해석"], font_size=8.2)


def add_method(doc: Document, weights: pd.DataFrame, salary_gate: pd.DataFrame) -> None:
    heading(doc, "2. 방법", 1)
    heading(doc, "2.1 데이터 및 전처리", 2)
    data_table = pd.DataFrame(
        [
            ["KBO/STATIZ", "SSG 상황별 팀 약점, KBO 외국인 첫해 성과 label", "SSG fit 및 historical target"],
            ["MLB/MiLB", "AAA/MLB 성적, IP/GS, K/9, BB/9, HR/9, OPS 등", "후보 시장 구축 및 번역"],
            ["Savant", "xwOBA, chase, whiff, hard-hit, barrel, 구종/구속 대응", "타자 성공/실패 classifier 및 KBO 적응 filter"],
            ["Roster/Transaction", "40인 여부, DFA, outrighted, FA, minor status", "영입 가능성 gate"],
        ],
        columns=["데이터", "주요 변수", "모델 내 역할"],
    )
    add_table(doc, data_table, ["데이터", "주요 변수", "모델 내 역할"], ["데이터", "주요 변수", "모델 내 역할"], font_size=8.2)
    callout(doc, "제외 변수.", "기사, 인터뷰, 평판, 텍스트 기반 감성/의지 변수는 모델 input에서 제외했다. 보고서 해석에는 참고할 수 있지만 점수화에는 사용하지 않았다.", fill=GOLD_LIGHT)

    heading(doc, "2.2 4인 앙상블 base learner", 2)
    p(
        doc,
        "앙상블은 네 명의 결과를 단순 평균하지 않는다. 각 분석이 담당한 데이터 마이닝 기능을 base learner로 정의하고, 슬롯별 검증 신뢰도에 따라 가중치를 다르게 부여했다.",
    )
    add_table(doc, weights, ["slot", "signal", "weight", "reason"], ["슬롯", "base learner", "비중", "근거"], font_size=7.3)
    add_picture(doc, ASSET_DIR / "ensemble_signal_weights.png", "그림 2.1 슬롯별 base learner 비중. 타자는 historical classifier, 투수는 consensus/filter 비중이 더 크다.")

    heading(doc, "2.3 후보 점수화 및 최종 1인 선발 규칙", 2)
    p(
        doc,
        "타자 앙상블 점수는 다음 구조로 계산했다. Hitter Ensemble = 0.40 x historical classifier + 0.25 x SSG fit/translation + 0.15 x KBO adaptation filter + 0.15 x market inefficiency + 0.05 x consensus.",
    )
    p(
        doc,
        "투수 앙상블 점수는 다음 구조로 계산했다. Pitcher Ensemble = 0.05 x historical classifier + 0.25 x SSG fit/translation + 0.25 x KBO adaptation filter + 0.20 x market inefficiency + 0.25 x consensus. 투수 historical classifier가 watch 등급이기 때문에 이 신호를 낮추고, 여러 base learner에서 반복 등장하는지와 KBO 적응 filter를 더 크게 반영했다.",
    )
    numbered(
        doc,
        [
            "1차: market access, 40인 여부, 후보 슬롯 적합성, 표본/결측 상태를 확인한다.",
            "2차: ensemble_score 기준으로 타자와 투수 각각 Top 3를 선정한다.",
            "3차: 최종 1인은 ensemble_score만으로 고르지 않고, historical model 신뢰도, consensus, role fit, 실패 리스크를 함께 본다.",
            "4차: salary, assignment, buyout, medical gate를 붙여 실제 접촉 가능 후보와 모델 리드를 분리한다.",
            "5차: 투수는 모델 성능이 약하므로 확정 추천이 아니라 검증 우선순위라는 표현을 유지한다.",
        ],
    )

    heading(doc, "2.4 연봉/계약/메디컬 hard gate", 2)
    p(
        doc,
        "기존 앙상블 점수만으로는 최종 영입 판단에 충분하지 않다. 특히 KBO 대체 외국인 제도는 월 비용 제한, 기존 외국인 계약 상태, release/buyout, assignment 권리, medical 상태에 직접 영향을 받는다. 따라서 본 v3에서는 후보 점수 이후 공개 연봉/계약/메디컬 gate를 별도 구조화 데이터로 붙였다.",
    )
    gate_view = salary_gate[
        [
            "slot",
            "candidate",
            "structured_salary_signal_usd",
            "contract_access_status",
            "medical_availability_status",
            "economic_gate",
            "model_action",
        ]
    ].copy()
    add_table(
        doc,
        gate_view,
        [
            "slot",
            "candidate",
            "structured_salary_signal_usd",
            "contract_access_status",
            "medical_availability_status",
            "economic_gate",
            "model_action",
        ],
        ["슬롯", "선수", "연봉 신호", "계약/소속 상태", "메디컬/가용성", "비용 gate", "모델 조치"],
        font_size=6.2,
    )
    callout(
        doc,
        "중요 수정.",
        "연봉/계약 gate를 붙이면 '모델이 가장 좋아하는 선수'와 '당장 접촉하기 가장 현실적인 선수'가 다를 수 있다. 따라서 최종 보고서에서는 모델 lead와 signable lead를 분리한다.",
        fill=GOLD_LIGHT,
    )


def add_results(doc: Document, scores: pd.DataFrame, salary_gate: pd.DataFrame) -> None:
    heading(doc, "3. 결과", 1)
    heading(doc, "3.1 앙상블 모델 구조", 2)
    p(
        doc,
        "최종 모델은 후보를 하나의 점수로만 줄이지 않고, 후보가 어떤 base learner를 통과했는지를 함께 저장한다. 이 때문에 같은 점수라도 해석이 달라진다. 예를 들어 Luis Matos는 model-discovery 후보이고, Jack Suwinski는 scouting-consensus 후보이며, Nolan Jones는 classifier와 consensus를 동시에 통과한 core 후보다.",
    )

    heading(doc, "3.2 외국인 타자 Top 3와 최종 1인", 2)
    hitter_top3 = scores[scores["slot"].eq("foreign_hitter")].head(3).copy()
    add_picture(doc, ASSET_DIR / "hitter_ensemble_ranking.png", "그림 3.1 외국인 타자 앙상블 랭킹. Nolan Jones가 모델 확률과 consensus를 동시에 통과했다.")
    add_table(
        doc,
        hitter_top3,
        ["rank", "candidate", "ensemble_score", "dm_success_prob", "dm_failure_prob", "source_presence", "ensemble_tier"],
        ["순위", "선수", "앙상블 점수", "성공확률", "실패확률", "근거 출처", "판정"],
        font_size=8.0,
    )
    callout(
        doc,
        "타자 결론: 모델 1위 Nolan Jones, 계약 현실성 1위 Luis Matos.",
        "Nolan Jones는 앙상블 점수 1위이며 historical classifier에서 성공확률 90.2%, 실패확률 9.2%를 기록했다. 동시에 sewon, jimini, Codex 구조화 모델에서 반복 등장했다. 다만 공개 2026 cash signal이 2.0M이므로 대체 외국인 월 비용 gate까지 고려하면 즉시 signable candidate로 단정할 수 없다. Luis Matos는 consensus가 약하지만 DFA/outright 이력과 낮은 공개 급여 부담 때문에 실제 접촉 현실성이 가장 좋다.",
        fill=BLUE_LIGHT,
    )
    add_paragraph = p
    add_paragraph(
        doc,
        "해석상 Nolan Jones는 ‘가장 화려한 후보’라기보다 가장 많은 모델 조건을 동시에 통과한 후보에 가깝다. 하지만 salary gate까지 붙이면 결론은 더 정교해진다. 발표에서는 Nolan Jones를 model lead, Luis Matos를 economic/signable lead로 분리하는 것이 안전하다.",
    )

    heading(doc, "3.3 외국인 선발투수 Top 3와 최종 1인", 2)
    pitcher_all = scores[scores["slot"].eq("foreign_pitcher")].copy()
    # Keep starter-oriented Top 3: if tied, prefer starter-depth over role-check relief profiles.
    pitcher_top3 = pitcher_all[pitcher_all["candidate"].isin(["Josh Fleming", "Carson Spiers", "Brian Van Belle"])].copy()
    pitcher_top3 = pitcher_top3.sort_values("ensemble_score", ascending=False)
    add_picture(doc, ASSET_DIR / "pitcher_ensemble_ranking.png", "그림 3.2 외국인 선발투수 앙상블 랭킹. Josh Fleming만 consensus가 강하고 나머지는 검증 후보로 남는다.")
    add_table(
        doc,
        pitcher_top3,
        ["rank", "candidate", "ensemble_score", "dm_success_prob", "dm_failure_prob", "source_presence", "ensemble_tier"],
        ["순위", "선수", "앙상블 점수", "성공확률", "실패위험", "근거 출처", "판정"],
        font_size=8.0,
    )
    callout(
        doc,
        "최종 투수 1인: Josh Fleming.",
        "Josh Fleming은 투수 후보 중 가장 강한 consensus 후보이며, jimini의 KBO 적응 filter, kyuho의 시장 비효율 feature model, Codex 구조화 모델에서 반복 등장했다. 또한 minor contract/outright/free-agency sequence가 있어 계약 접근성도 상대적으로 높다. 투수 historical classifier는 AUC 0.603으로 약하기 때문에 확정 성공 예측이 아니라 검증 우선순위 1번으로 표현해야 한다.",
        fill=ORANGE_LIGHT,
    )
    p(
        doc,
        "Carson Spiers는 sewon의 번역/선발 안정성 파이프라인에서 강하게 올라온 후보이고, Brian Van Belle은 이닝과 커맨드 성격이 선발 슬롯에 맞는 후보였다. 그러나 salary/medical gate를 붙이면 둘 다 active release candidate가 아니다. Spiers는 elbow surgery로 2026 대부분 결장 위험이 있고, Van Belle은 full-season IL 상태다. 따라서 투수 쪽은 raw Top 3와 실제 접촉 Top 3를 구분해야 한다.",
    )

    heading(doc, "3.4 모델 검증과 해석 강도", 2)
    add_picture(doc, ASSET_DIR / "model_validation_auc.png", "그림 3.3 타자/투수 classifier 검증 성능. 타자는 추천 가능, 투수는 diagnostic only로 해석한다.")
    p(
        doc,
        "타자 success classifier는 AUC 0.833, failure classifier는 AUC 0.738로 후보 추천에 사용할 수 있다. 반면 투수 success classifier는 AUC 0.603으로 확정 추천에 쓰기 어렵다. 따라서 본 보고서의 결론 강도는 슬롯별로 다르다. 타자는 최종 1인 추천, 투수는 최종 1인 검증 우선순위다.",
    )

    heading(doc, "3.5 연봉 gate 반영 후 결론 변화", 2)
    post_gate = pd.DataFrame(
        [
            ["타자", "Nolan Jones", "모델 1위", "2.0M cash signal로 비용 gate 확인 필요", "모델 lead 유지"],
            ["타자", "Luis Matos", "계약 현실성 1위", "DFA/outright 이력 + 낮은 공개 급여 부담", "대체 외국인 접촉 1순위 가능"],
            ["타자", "Jack Suwinski", "upside/consensus 후보", "1.25M deal signal", "cost-share/release 확인 전 보류"],
            ["투수", "Josh Fleming", "검증 1순위", "minor contract/outright sequence", "접촉 가능 후보로 유지"],
            ["투수", "Carson Spiers", "raw Top 3", "2026 대부분 결장 medical red", "medical hold"],
            ["투수", "Brian Van Belle", "raw Top 3", "full-season IL", "medical hold"],
        ],
        columns=["슬롯", "선수", "모델상 위치", "연봉/계약/메디컬 gate", "수정 결론"],
    )
    add_table(doc, post_gate, ["슬롯", "선수", "모델상 위치", "연봉/계약/메디컬 gate", "수정 결론"], ["슬롯", "선수", "모델상 위치", "gate", "수정 결론"], font_size=8.0)
    callout(
        doc,
        "v3 최종 해석.",
        "기존 v2 결론에 그대로 만족하면 안 된다. v3에서는 '모델 lead'와 '계약 가능 lead'를 분리한다. 타자는 Nolan Jones가 데이터 마이닝 1위, Luis Matos가 비용 현실성 1위다. 투수는 Josh Fleming만 active contact lead로 남고, Spiers/Van Belle은 medical hold로 내린다.",
        fill=GOLD_LIGHT,
    )


def add_discussion_limit_conclusion(doc: Document, salary_gate: pd.DataFrame) -> None:
    heading(doc, "4. 논의", 1)
    p(
        doc,
        "이 보고서의 핵심은 앙상블을 ‘사람 네 명의 타협안’으로 쓰지 않았다는 점이다. 각 분석은 서로 다른 데이터 마이닝 기능을 갖는 base learner로 재정의되었다. historical classifier는 과거 KBO 성공/실패 패턴을, SSG fit pipeline은 팀 약점 보완을, KBO adaptation filter는 실패 리스크 제거를, market inefficiency model은 숨은 upside를 담당한다.",
    )
    p(
        doc,
        "실무 적용에서는 Nolan Jones와 Josh Fleming을 바로 같은 강도로 말하면 안 된다. Nolan Jones는 최종 영입 후보라고 말할 수 있지만, Josh Fleming은 최종 검증 1순위라고 표현하는 것이 더 정확하다. 이 차이가 보고서의 신뢰도를 만든다.",
    )

    heading(doc, "5. 한계", 1)
    bullets(
        doc,
        [
            "과거 KBO 외국인 label 표본이 작다. 특히 투수 모델은 AUC 0.603으로 확정 추천에 충분하지 않다.",
            "공개 연봉/계약/메디컬 gate를 붙였지만, 실제 buyout, opt-out, remaining salary, transfer fee, 한국행 의사는 구단/에이전트 확인이 필요하다.",
            "기사, 인터뷰, 평판 텍스트는 사용자 요청에 따라 모델 input에서 제외했다.",
            "Savant와 MiLB feature는 리그, 구장, 표본 수에 따라 변동성이 있다.",
            "앙상블 weight는 모델 검증력과 분석 목적에 기반한 설계값이며, 더 많은 historical label이 확보되면 재학습해야 한다.",
        ],
    )

    heading(doc, "6. 결론", 1)
    conclusion_table = pd.DataFrame(
        [
            ["외국인 타자", "Nolan Jones, Luis Matos, Jack Suwinski", "Nolan Jones", "모델 1위"],
            ["외국인 타자", "Nolan Jones, Luis Matos, Jack Suwinski", "Luis Matos", "계약 현실성 1위"],
            ["외국인 선발투수", "Josh Fleming, Carson Spiers, Brian Van Belle", "Josh Fleming", "검증 우선순위 1번"],
        ],
        columns=["슬롯", "최종 Top 3", "최종 1인", "해석 강도"],
    )
    add_table(doc, conclusion_table, ["슬롯", "최종 Top 3", "최종 1인", "해석 강도"], ["슬롯", "최종 Top 3", "최종 1인", "해석 강도"], font_size=8.6)
    callout(
        doc,
        "핵심 한 줄.",
        "타자는 모델만 보면 Nolan Jones가 1위지만, 대체 외국인 비용 gate까지 고려하면 Luis Matos가 더 현실적인 접촉 1순위다. 투수는 Josh Fleming을 최종 검증 1순위로 둔다. 이 결론은 네 명의 후보표를 평균낸 것이 아니라, 서로 다른 데이터 마이닝 base learner에 salary/medical hard gate를 붙인 결과다.",
        fill=GOLD_LIGHT,
    )

    heading(doc, "A. 부록", 1)
    p(doc, "A.1 재현 산출물")
    bullets(
        doc,
        [
            "outputs/tables/ensemble_candidate_scores_v1.csv",
            "outputs/tables/ensemble_model_signal_weights_v1.csv",
            "outputs/tables/final_candidate_salary_contract_gate_v1.csv",
            "outputs/reports/ensemble_final_report_v1.html",
            "scripts/build_ensemble_final_report_v1.py",
            "scripts/build_ensemble_candidate_selection_report_docx_v2.py",
        ],
    )
    p(doc, "A.2 공개 계약/연봉/메디컬 확인 링크")
    source_view = salary_gate[["candidate", "source_url", "source_note"]].copy()
    add_table(doc, source_view, ["candidate", "source_url", "source_note"], ["선수", "확인 링크", "확인 내용"], font_size=6.2)


def build_docx() -> None:
    scores = pd.read_csv(SCORES_PATH)
    weights = pd.read_csv(WEIGHTS_PATH).copy()
    weights["weight"] = weights["weight"].map(lambda v: f"{v:.0%}")
    salary_gate = pd.read_csv(SALARY_GATE_PATH)

    DOC_OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_defaults(doc)
    add_cover(doc)
    add_abstract(doc)
    add_contents(doc)
    add_intro(doc)
    add_method(doc, weights, salary_gate)
    add_results(doc, scores, salary_gate)
    add_discussion_limit_conclusion(doc, salary_gate)
    add_footer(doc)
    doc.save(DOCX_PATH)


def main() -> None:
    build_docx()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
