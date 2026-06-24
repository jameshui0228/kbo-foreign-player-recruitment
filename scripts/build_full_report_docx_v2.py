#!/usr/bin/env python3
"""Build a full Korean DOCX report for the SSG foreign-player recruitment project."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "tables"
DOC_OUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = ROOT / "tmp" / "docs" / "ssg_full_report_v2_assets"
DOC_PATH = DOC_OUT_DIR / "SSG_외국인선수_데이터마이닝_정식보고서_v2.docx"


def read_csv(name: str) -> pd.DataFrame:
    path = OUT / name
    if not path.exists():
        raise FileNotFoundError(path)
    return pd.read_csv(path)


def fmt_pct(value: float | int | None, digits: int = 1) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{float(value) * 100:.{digits}f}%"


def fmt_num(value: float | int | None, digits: int = 3) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{float(value):.{digits}f}"


def fmt_int(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{int(round(float(value))):,}"


def fmt_money(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "미확인"
    return f"${int(round(float(value))):,}"


def set_run_font(run, size: int = 10, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: object,
    *,
    bold: bool = False,
    size: int = 8,
    color: str = "111827",
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run("" if text is None else str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths: list[float] | None) -> None:
    if not widths:
        return
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Cm(width)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 12, bold=True, color="111827")


def add_para(document: Document, text: str, *, size: int = 10, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=size, bold=True)
        rest = p.add_run(text[len(bold_prefix) :])
        set_run_font(rest, size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size)


def add_bullet(document: Document, text: str, *, size: int = 9) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size)


def add_numbered(document: Document, text: str, *, size: int = 9) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size)


def add_callout(document: Document, title: str, body: str, *, fill: str = "EAF2FF") -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_run_font(run, size=10, bold=True)
    p.add_run("\n")
    for run in p.runs[1:]:
        set_run_font(run, size=9)
    body_run = p.add_run(body)
    set_run_font(body_run, size=9)
    document.add_paragraph()


def add_table(
    document: Document,
    headers: list[str],
    rows: Iterable[Iterable[object]],
    *,
    widths: list[float] | None = None,
    font_size: int = 7,
    header_fill: str = "1F2937",
) -> None:
    rows = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, size=font_size, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    set_table_widths(table, widths)
    document.add_paragraph()


def add_picture(document: Document, image_path: Path, width_inches: float = 6.4) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))


def apply_doc_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    styles["Normal"].font.size = Pt(9.5)


def make_charts() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.family"] = ["AppleGothic", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    charts: dict[str, Path] = {}

    runway = read_csv("ssg_2026_runway_gap_by_team.csv")
    ssg = runway[runway["t_code_name"].eq("SSG")].iloc[0]
    fig, ax = plt.subplots(figsize=(7.2, 3.5))
    labels = ["RISP OPS\n(KBO rank 1)", "Runner on 1B OPS\n(KBO rank 10)", "Runner on 1B OBP\n(KBO rank 10)"]
    values = [ssg["OPS_risp"], ssg["OPS_on_first"], ssg["OBP_on_first"]]
    colors = ["#2563EB", "#DC2626", "#F97316"]
    ax.bar(labels, values, color=colors)
    ax.set_ylim(0, 0.95)
    ax.set_ylabel("Rate")
    ax.set_title("SSG 2026: RISP production exists, first-base runway collapses")
    for i, value in enumerate(values):
        ax.text(i, value + 0.025, f"{value:.3f}", ha="center", fontsize=10, fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    path = ASSET_DIR / "runway_gap.png"
    fig.savefig(path, dpi=200)
    plt.close(fig)
    charts["runway_gap"] = path

    pitch_ctx = read_csv("ssg_pitching_message_v0_2_context_validation.csv").head(6)
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.barh(pitch_ctx["context_label"][::-1], pitch_ctx["bad_context_score"][::-1], color="#B91C1C")
    ax.set_xlim(0, 10.5)
    ax.set_xlabel("Bad-context score")
    ax.set_title("SSG pitching worst contexts: traffic, early innings, right-handed lineups")
    for i, value in enumerate(pitch_ctx["bad_context_score"][::-1]):
        ax.text(value + 0.1, i, f"{value:.1f}", va="center", fontsize=9)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    path = ASSET_DIR / "pitching_context.png"
    fig.savefig(path, dpi=200)
    plt.close(fig)
    charts["pitching_context"] = path

    decisions = read_csv("kbo_translation_failure_feature_family_decisions_v0_3.csv")
    model_rows = decisions[
        (decisions["role_model_family"].isin(["hitter", "pitcher"]))
        & (decisions["target"].isin(["success", "failure"]))
    ].copy()
    model_rows["label"] = model_rows["role_model_family"] + " " + model_rows["target"]
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    colors = ["#2563EB" if x == "hitter" else "#6B7280" for x in model_rows["role_model_family"]]
    ax.bar(model_rows["label"], model_rows["mean_auc"], color=colors)
    ax.axhline(0.5, color="#111827", linewidth=1, linestyle="--")
    ax.set_ylim(0.45, 0.9)
    ax.set_ylabel("Repeated CV AUC")
    ax.set_title("Model validation: hitter classifiers are usable, pitcher is diagnostic")
    ax.tick_params(axis="x", rotation=20)
    for i, value in enumerate(model_rows["mean_auc"]):
        ax.text(i, value + 0.012, f"{value:.3f}", ha="center", fontsize=9)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    path = ASSET_DIR / "model_validation_auc.png"
    fig.savefig(path, dpi=200)
    plt.close(fig)
    charts["model_validation_auc"] = path

    top = read_csv("data_mining_recommendations_top3_v1.csv")
    fig, ax = plt.subplots(figsize=(7.2, 4.4))
    names = top["player_name"].tolist()[::-1]
    success = top["dm_success_prob"].tolist()[::-1]
    failure = top["dm_failure_prob"].tolist()[::-1]
    y = range(len(names))
    ax.barh(y, success, color="#16A34A", label="P(success)")
    ax.barh(y, [-x for x in failure], color="#EF4444", label="P(failure)")
    ax.axvline(0, color="#111827", linewidth=0.8)
    ax.set_yticks(list(y))
    ax.set_yticklabels(names)
    ax.set_xlim(-0.65, 1.0)
    ax.set_xlabel("Probability")
    ax.set_title("Data-mining top candidates: success vs failure signal")
    ax.legend(loc="lower right", fontsize=8)
    for i, (s, f) in enumerate(zip(success, failure)):
        ax.text(s + 0.02, i, f"{s:.2f}", va="center", fontsize=8)
        ax.text(-f - 0.08, i, f"{f:.2f}", va="center", fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    path = ASSET_DIR / "candidate_probabilities.png"
    fig.savefig(path, dpi=200)
    plt.close(fig)
    charts["candidate_probabilities"] = path

    return charts


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("SSG 랜더스 대체 외국인/아시아쿼터\n선수 영입 데이터 마이닝 정식 보고서")
    set_run_font(run, size=22, bold=True, color="111827")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SDA 2기 스포츠 데이터 분석 학술 동아리 | v2 | 2026-06-22 KST")
    set_run_font(run, size=10, color="4B5563")

    add_callout(
        document,
        "보고서의 핵심",
        "이번 문서는 단순 후보명 정리가 아니라, 1) SSG 숨은 약점, 2) KBO 외인 성공/실패 유형, "
        "3) 후보 시장, 4) KBO 번역 모델, 5) 실패 리스크, 6) 최종 SSG fit ranking이 어떻게 하나의 "
        "데이터 마이닝 결론으로 연결되는지 설명한다. 기사/인터뷰/뉴스 텍스트는 최종 모델 입력에서 제외했고, "
        "숫자형 성적 및 구조화 데이터만 사용했다.",
        fill="DDEBFF",
    )

    add_heading(document, "한 줄 결론", 1)
    add_bullet(document, "외인타자: 모델이 강하게 지지하는 후보는 Luis Matos, Nolan Jones, Dylan Carlson이다.")
    add_bullet(document, "외인투수: Bryse Wilson과 Austin Gomber는 진단 리드지만, 투수 모델은 아직 확정 추천이라고 말하면 과장이다.")
    add_bullet(document, "SSG 메시지: 타자는 순수 장타가 아니라 first-base traffic converter, 투수는 탈삼진 쇼케이스가 아니라 load-bearing traffic-command starter가 핵심이다.")
    add_bullet(document, "현실성: 40인 로스터 제외 여부는 반영했지만, 최종 계약 가능성은 공개 연봉/opt-out/보장계약/메디컬 확인을 추가로 거쳐야 한다.")

    document.add_page_break()


def add_table_of_contents(document: Document) -> None:
    add_heading(document, "목차", 1)
    for item in [
        "1. 데이터와 모델 설계",
        "2. SSG 숨은 약점 마이닝: 타격",
        "3. SSG 숨은 약점 마이닝: 투수",
        "4. KBO 외국인 성공/실패 유형 마이닝",
        "5. 후보 시장 구축 및 계약/규정 현실성",
        "6. KBO 번역 모델과 실패 리스크 모델",
        "7. 최종 후보 및 후보별 통과 논리",
        "8. 이전 점수표와 데이터 마이닝 결론의 차이",
        "9. 최종 메시지와 다음 액션",
        "10. Appendix: 사용 산출물",
    ]:
        add_numbered(document, item, size=10)
    document.add_page_break()


def add_data_and_model_design(document: Document, charts: dict[str, Path]) -> None:
    add_heading(document, "1. 데이터와 모델 설계", 1)
    add_callout(
        document,
        "모델링 원칙",
        "사용자가 지정한 기준에 맞춰 기사, 인터뷰, 뉴스, 글 기반 변수는 최종 모델 input에서 제외했다. "
        "뉴스/인터뷰는 배경 이해에는 쓸 수 있지만, 최종 후보 산출 모델에는 숫자형 성적 데이터와 구조화 데이터만 들어간다.",
        fill="E8F5E9",
    )

    coverage = read_csv("data_coverage_by_league_v1.csv")
    wanted = [
        "KBO/STATIZ snapshot",
        "MLB Savant raw 2023",
        "MLB Savant raw 2024",
        "MLB Savant raw 2025",
        "MLB Savant raw 2026",
        "MLB roster status latest",
        "MiLB current/historical output",
        "NPB official roster output",
        "NPB official stats output",
        "CPBL official roster output",
        "NPB+CPBL Asian quota market output",
    ]
    table_rows = []
    for _, row in coverage[coverage["league_bucket"].isin(wanted)].iterrows():
        table_rows.append(
            [
                row["league_bucket"],
                fmt_int(row["files"]),
                fmt_int(row["rows"]),
                f"{float(row['size_mb']):.1f}MB" if pd.notna(row["size_mb"]) else "-",
                row["detail"],
            ]
        )
    add_table(
        document,
        ["데이터 묶음", "파일", "행", "크기", "설명"],
        table_rows,
        widths=[4.2, 1.1, 2.0, 1.7, 6.2],
        font_size=7,
    )

    audit = read_csv("project_data_coverage_audit_v1.csv")
    audit_rows = []
    for _, row in audit.head(5).iterrows():
        audit_rows.append([row["source_group"], fmt_int(row["rows"]), row["coverage"], row["candidate_use"], row["status"]])
    add_table(
        document,
        ["Source", "Rows", "Coverage", "Candidate use", "Status"],
        audit_rows,
        widths=[3.6, 1.8, 4.1, 5.0, 2.0],
        font_size=6,
    )

    progress = read_csv("recruitment_gate_status_v33.csv")
    progress_rows = []
    for _, row in progress.iterrows():
        progress_rows.append([row["gate"], row["layer"], f"{int(row['progress_pct'])}%", row["status"], row["blocking_gap"]])
    add_table(
        document,
        ["Gate", "Layer", "진행률", "현재 상태", "남은 gap"],
        progress_rows,
        widths=[1.3, 4.2, 1.3, 4.0, 5.2],
        font_size=6,
    )

    add_heading(document, "통합 모델의 6단계", 2)
    for item in [
        "SSG 숨은 약점 마이닝: SSG가 진짜로 보강해야 하는 game-state 병목을 찾는다.",
        "KBO 외인 성공/실패 유형 마이닝: 과거 KBO 외인의 성공/실패 archetype을 학습한다.",
        "후보 시장 구축: MLB/AAA/AA/NPB/CPBL 후보를 로스터, 포지션, 시장 접근성 기준으로 구성한다.",
        "KBO 번역 모델: 해외 성과가 KBO에서 성공/실패로 번역될 가능성을 추정한다.",
        "실패 리스크 모델: 메디컬, 계약, 로스터, translation uncertainty를 구조화 penalty로 반영한다.",
        "SSG fit ranking: 위 블록을 하나로 합쳐 외인타자/외인투수 후보를 뽑는다.",
    ]:
        add_bullet(document, item)

    add_picture(document, charts["model_validation_auc"], width_inches=6.5)
    add_callout(
        document,
        "중요한 방어 논리",
        "타자 모델은 repeated CV에서 성공 AUC 0.833, 실패 AUC 0.738로 pilot_promote를 통과했다. "
        "반면 투수 모델은 success AUC 0.603의 watch 수준이라, 투수 후보는 확정 추천이 아니라 진단 리드로 발표하는 것이 안전하다.",
        fill="FFF7D6",
    )
    document.add_page_break()


def add_ssg_hitting_layer(document: Document, charts: dict[str, Path]) -> None:
    add_heading(document, "2. SSG 숨은 약점 마이닝: 타격", 1)
    add_picture(document, charts["runway_gap"], width_inches=6.5)

    runway = read_csv("ssg_2026_runway_gap_by_team.csv")
    ssg = runway[runway["t_code_name"].eq("SSG")].iloc[0]
    add_callout(
        document,
        "핵심 발견",
        f"SSG는 2026년 득점권 OPS가 {ssg['OPS_risp']:.3f}로 KBO 1위인데, 1루 주자 상황 OPS는 "
        f"{ssg['OPS_on_first']:.3f}로 10위다. 즉 문제는 '득점권에서 못 친다'가 아니라, 1루 주자를 "
        "득점권으로 옮기는 runway가 끊기는 쪽에 있다.",
        fill="EAF2FF",
    )

    runway_rows = []
    for _, row in runway.head(5).iterrows():
        runway_rows.append(
            [
                row["t_code_name"],
                fmt_num(row["OPS_risp"], 3),
                fmt_int(row["risp_ops_rank"]),
                fmt_num(row["OPS_on_first"], 3),
                fmt_int(row["on_first_ops_rank"]),
                fmt_num(row["OBP_on_first"], 3),
                fmt_num(row["risp_minus_on_first_ops"], 3),
            ]
        )
    add_table(
        document,
        ["팀", "RISP OPS", "RISP 순위", "1루 OPS", "1루 순위", "1루 OBP", "Gap"],
        runway_rows,
        widths=[1.6, 1.8, 1.4, 1.8, 1.4, 1.8, 1.6],
        font_size=7,
    )

    role = read_csv("ssg_2026_role_runway_context.csv")
    role_focus = role[role["split_label"].isin(["on_first", "lt2_out_on_first", "risp"])].copy()
    show_segments = ["IF_C_core", "OF_high_leverage_usage", "OF_lower_or_mixed_usage", "DH_primary_or_bridge"]
    role_focus = role_focus[role_focus["role_segment"].isin(show_segments)]
    role_rows = []
    for _, row in role_focus.iterrows():
        role_rows.append(
            [
                row["split_label"],
                row["role_segment"],
                fmt_int(row["pa"]),
                fmt_num(row["obp"], 3),
                fmt_num(row["slg"], 3),
                fmt_num(row["ops"], 3),
                fmt_pct(row["gdp_per_pa"], 1),
                fmt_pct(row["k_pct"], 1),
            ]
        )
    add_table(
        document,
        ["상황", "Role segment", "PA", "OBP", "SLG", "OPS", "GDP/PA", "K%"],
        role_rows,
        widths=[2.3, 3.8, 1.2, 1.4, 1.4, 1.4, 1.5, 1.4],
        font_size=6,
    )

    robust = read_csv("ssg_hidden_weakness_robustness_decisions_v3.csv")
    robust_rows = []
    for _, row in robust.iterrows():
        robust_rows.append(
            [
                row["rule_id"],
                row["core_type"],
                fmt_int(row["games"]),
                fmt_pct(row["win_pct"], 1),
                fmt_num(row["avg_run_diff"], 2),
                fmt_pct(row["leave_one_opponent_support_rate"], 1),
                row["decision"],
            ]
        )
    add_table(
        document,
        ["Rule", "의미", "G", "승률", "평균 득실", "상대팀 제거 안정성", "판정"],
        robust_rows,
        widths=[4.4, 2.8, 0.9, 1.2, 1.4, 2.0, 2.0],
        font_size=6,
    )

    add_heading(document, "타자 보강 메시지", 2)
    for item in [
        "SSG의 타자 보강 메시지는 '장타 부족'이 아니다. 이미 득점권 생산성은 강하다.",
        "핵심은 우투수 상대 OF/DH 구간에서 1루 주자 상황을 살리는 능력이다.",
        "외인타자는 first-base traffic converter: 출루, 라인드라이브/갭 타구, two-strike survival, 병살 회피, 우투수 대응이 필요하다.",
        "따라서 단순 홈런형보다 '안 죽고 다음 상황을 만드는 외야수'가 SSG 맞춤형이다.",
    ]:
        add_bullet(document, item)
    document.add_page_break()


def add_ssg_pitching_layer(document: Document, charts: dict[str, Path]) -> None:
    add_heading(document, "3. SSG 숨은 약점 마이닝: 투수", 1)
    add_picture(document, charts["pitching_context"], width_inches=6.5)

    decision = read_csv("ssg_pitching_message_v0_2_decision_table.csv")
    decision_rows = []
    for _, row in decision.iterrows():
        decision_rows.append([row["criterion"], fmt_num(row["value"], 3), row["threshold"], row["status"], row["note"]])
    add_table(
        document,
        ["Criterion", "Value", "Threshold", "Status", "Note"],
        decision_rows,
        widths=[4.1, 1.5, 2.7, 1.3, 5.8],
        font_size=6,
    )

    context = read_csv("ssg_pitching_message_v0_2_context_validation.csv").head(8)
    context_rows = []
    for _, row in context.iterrows():
        context_rows.append(
            [
                row["context_label"],
                fmt_int(row["G"]),
                fmt_num(row["ERA"], 2),
                fmt_num(row["WHIP"], 2),
                fmt_num(row["OPS"], 3),
                fmt_num(row["bb9"], 2),
                fmt_num(row["hr9"], 2),
                fmt_num(row["bad_context_score"], 2),
            ]
        )
    add_table(
        document,
        ["Context", "G", "ERA", "WHIP", "OPS", "BB/9", "HR/9", "Bad score"],
        context_rows,
        widths=[3.7, 1.0, 1.3, 1.3, 1.3, 1.3, 1.3, 1.7],
        font_size=7,
    )

    impact = read_csv("ssg_2026_import_slot_pitching_impact.csv")
    impact_rows = []
    for _, row in impact.iterrows():
        impact_rows.append(
            [
                row["import_slot_group"],
                row["pitch_role"],
                fmt_int(row["starts"]),
                fmt_num(row["ip_per_game"], 2),
                fmt_num(row["era"], 2),
                fmt_num(row["whip"], 2),
                fmt_num(row["bb9"], 2),
                fmt_num(row["ops_allowed"], 3),
            ]
        )
    add_table(
        document,
        ["Group", "Role", "Starts", "IP/G", "ERA", "WHIP", "BB/9", "OPS allowed"],
        impact_rows,
        widths=[3.0, 1.7, 1.2, 1.3, 1.3, 1.3, 1.3, 1.8],
        font_size=7,
    )

    add_callout(
        document,
        "핵심 발견",
        "SSG 외국인 투수 슬롯은 팀의 scarcity를 흡수하지 못하고 오히려 증폭하고 있다. "
        "import-slot starter는 ERA 6.17, WHIP 1.73, 피OPS .821로 domestic starter보다 더 나쁜 run prevention을 보였고, "
        "그 결과 SSG가 따라갈 경기의 runway가 초반에 끊긴다.",
        fill="FFE4E6",
    )

    add_heading(document, "투수 보강 메시지", 2)
    for item in [
        "SSG에 필요한 외인투수는 단순 구속/탈삼진형이 아니라 traffic-command starter다.",
        "목표는 5이닝 초반을 6이닝으로 바꾸고, 초반 볼넷/피홈런/장타 허용을 줄이는 것이다.",
        "KBO 성공 외인투수 archetype도 '많은 이닝을 안정적으로 먹는 rotation anchor'에 가깝다.",
        "따라서 후보 평가에서 K/9만 크게 보면 안 되고 BB/9, HR/9, WHIP, workload continuity, runner-on-base damage control을 같이 봐야 한다.",
    ]:
        add_bullet(document, item)
    document.add_page_break()


def add_kbo_archetype_layer(document: Document) -> None:
    add_heading(document, "4. KBO 외국인 성공/실패 유형 마이닝", 1)
    arche = read_csv("kbo_foreign_archetype_prearrival_profile_v0_2.csv")

    arche_rows = []
    keep = [
        "hitter_failure_replacement_or_low_impact",
        "hitter_everyday_middle_order_anchor",
        "pitcher_load_bearing_rotation_anchor",
        "pitcher_failure_replacement_or_health_risk",
    ]
    for _, row in arche[arche["archetype_name"].isin(keep)].iterrows():
        arche_rows.append(
            [
                row["role_model_family"],
                row["archetype_name"],
                fmt_int(row["rows"]),
                fmt_pct(row["success_rate"], 1),
                fmt_pct(row["failure_rate"], 1),
                fmt_pct(row["in_season_replaced_rate"], 1),
                fmt_num(row["median_first_kbo_war"], 2),
            ]
        )
    add_table(
        document,
        ["Role", "Archetype", "Rows", "Success", "Failure", "In-season replaced", "Median WAR"],
        arche_rows,
        widths=[1.5, 5.7, 1.0, 1.5, 1.5, 2.2, 1.5],
        font_size=6,
    )

    failure = read_csv("kbo_foreign_failure_patterns_v0_1.csv")
    failure_rows = []
    for _, row in failure.iterrows():
        failure_rows.append(
            [
                row["role_group"],
                row["failure_archetype"],
                fmt_int(row["failure_rows"]),
                fmt_pct(row["share_of_role_failures"], 1),
                fmt_pct(row["injury_exit_rate"], 1),
                fmt_pct(row["performance_exit_rate"], 1),
                fmt_pct(row["in_season_replaced_rate"], 1),
                fmt_num(row["median_war"], 2),
            ]
        )
    add_table(
        document,
        ["Role", "Failure archetype", "Rows", "Share", "Injury", "Performance", "Replaced", "Median WAR"],
        failure_rows,
        widths=[1.5, 5.7, 1.0, 1.3, 1.3, 1.8, 1.5, 1.4],
        font_size=6,
    )

    add_callout(
        document,
        "모델 메시지",
        "KBO 외국인 성공은 '가장 화려한 툴'보다 '한 시즌을 버티는 역할 안정성'과 더 강하게 연결된다. "
        "타자는 everyday middle-order anchor, 투수는 load-bearing rotation anchor가 성공 archetype이다. "
        "반대로 실패는 낮은 출전량, 시즌 중 교체, 부상 이탈, performance exit에 집중된다.",
        fill="EAF2FF",
    )
    document.add_page_break()


def add_market_and_contract_layer(document: Document) -> None:
    add_heading(document, "5. 후보 시장 구축 및 계약/규정 현실성", 1)
    market = read_csv("candidate_market_coverage_v0_3.csv")
    market_rows = []
    for _, row in market[market["rows"].fillna(0).gt(0)].iterrows():
        market_rows.append(
            [
                row["slot"],
                fmt_int(row["rows"]),
                fmt_int(row["research_lead_rows"]),
                fmt_int(row["market_watch_rows"]),
                fmt_int(row["medical_hold_rows"]),
                fmt_int(row["recent_release_or_dfa_rows"]),
                row["blocking_gap"],
            ]
        )
    add_table(
        document,
        ["Slot", "Rows", "Research", "Watch", "Medical hold", "Recent DFA/release", "Blocking gap"],
        market_rows,
        widths=[3.7, 1.0, 1.3, 1.3, 1.7, 1.9, 5.7],
        font_size=6,
    )

    risk = read_csv("candidate_failure_risk_slot_summary_v0_1.csv")
    risk_focus = risk[risk["fit_slot"].isin(["foreign_hitter", "foreign_pitcher", "asian_quota"])].copy()
    risk_rows = []
    for _, row in risk_focus.iterrows():
        risk_rows.append(
            [
                row["fit_slot"],
                row["failure_risk_review_tier"],
                fmt_int(row["rows"]),
                fmt_num(row["median_failure_risk"], 1),
                fmt_int(row["high_medical_rows"]),
                fmt_int(row["high_contract_rows"]),
                str(row["release_allowed"]),
            ]
        )
    add_table(
        document,
        ["Fit slot", "Risk tier", "Rows", "Median risk", "Medical", "Contract", "Release allowed"],
        risk_rows,
        widths=[2.2, 4.1, 1.0, 1.6, 1.4, 1.4, 1.9],
        font_size=6,
    )

    constraints = read_csv("kbo_contract_constraints_v1.csv")
    constraint_rows = []
    for _, row in constraints.head(8).iterrows():
        constraint_rows.append([row["market_bucket"], row["constraint_name"], row["rule_value"], row["model_action"], row["status"]])
    add_table(
        document,
        ["Market", "Constraint", "Rule", "Model action", "Status"],
        constraint_rows,
        widths=[2.0, 3.0, 7.2, 2.0, 1.5],
        font_size=6,
    )

    salary_rows = [
        ["Luis Matos", "외인타자", "$129k-$780k range", "Spotrac payroll/ranking range; exact movable cost 확인 필요", "저비용 가능성은 있으나 조직 통제/계약 의사 확인 필요"],
        ["Nolan Jones", "외인타자", "$2.0M", "Spotrac/Guardians payroll; 2026 보장액 신호", "비용/이적 가능성 gate가 가장 중요"],
        ["Dylan Carlson", "외인타자", "$2.0M MLB roster signal / minor terms unclear", "2026 계약 조건 소스 불일치", "현재 조직/보장액 재확인 전 확정 금지"],
        ["Bryse Wilson", "외인투수", "$850k signal", "Spotrac/Cot's Phillies salary signal", "진단 리드이나 salary/role 확인 필요"],
        ["Austin Gomber", "외인투수", "undisclosed minor deal", "MiLB transaction: Atlanta/Gwinnett minor-league context", "접근성은 있으나 정확 비용 미확인"],
        ["Dietrich Enns", "외인투수", "$2.5M-$2.625M signal", "Orioles/Spotrac salary signal", "모델 margin 음수 + 비용 block 성격"],
    ]
    add_heading(document, "공개 연봉/계약 현실성 게이트", 2)
    add_table(
        document,
        ["선수", "Slot", "공개 비용 신호", "근거", "해석"],
        salary_rows,
        widths=[2.6, 1.6, 3.0, 4.4, 4.4],
        font_size=6,
    )
    add_callout(
        document,
        "연봉 데이터의 위치",
        "현재 최종 데이터 마이닝 모델은 선수 성과와 구조화 시장 접근성을 중심으로 작동한다. "
        "공개 연봉은 아직 전 후보군 전체에 완전하게 붙지 않았기 때문에 모델의 핵심 input이 아니라 최종 현실성 gate로 둔다. "
        "다음 단계는 Top 20 후보 전체에 current salary, guarantee, opt-out, assignment, buyout을 붙여 hard gate로 만드는 것이다.",
        fill="FFF7D6",
    )
    document.add_page_break()


def add_translation_and_risk_layer(document: Document) -> None:
    add_heading(document, "6. KBO 번역 모델과 실패 리스크 모델", 1)
    readiness = read_csv("kbo_translation_model_readiness_v0_3.csv")
    ready_rows = []
    for _, row in readiness.iterrows():
        ready_rows.append(
            [
                row["scope"],
                fmt_int(row["rows"]),
                fmt_int(row["trainable_rows"]),
                fmt_pct(row["model_ready_rate"], 1),
                fmt_int(row["repeated_cv_rows"]),
                row["release_policy"],
            ]
        )
    add_table(
        document,
        ["Scope", "Rows", "Trainable", "Ready rate", "CV rows", "Policy"],
        ready_rows,
        widths=[1.6, 1.2, 1.4, 1.5, 1.3, 7.0],
        font_size=6,
    )

    decisions = read_csv("kbo_translation_failure_feature_family_decisions_v0_3.csv")
    decision_rows = []
    for _, row in decisions.iterrows():
        decision_rows.append(
            [
                row["role_model_family"],
                row["target"],
                row["feature_family"],
                row["model"],
                fmt_num(row["mean_auc"], 3),
                fmt_num(row["mean_brier"], 3),
                fmt_num(row["brier_lift_vs_role_prior"], 3),
                fmt_num(row["top25_precision_lift_vs_role_prior"], 3),
                row["promotion_status"],
            ]
        )
    add_table(
        document,
        ["Role", "Target", "Feature", "Model", "AUC", "Brier", "Brier lift", "Top25 lift", "Status"],
        decision_rows,
        widths=[1.5, 1.4, 3.0, 2.4, 1.1, 1.2, 1.4, 1.5, 2.0],
        font_size=6,
    )

    locked = read_csv("layer5_failure_risk_v0_3_slot_summary_v0_1.csv")
    locked_rows = []
    for _, row in locked.iterrows():
        locked_rows.append(
            [
                row["fit_slot"],
                row["failure_risk_band_v0_3"],
                row["manual_review_tier_v0_3"],
                fmt_int(row["locked_rows"]),
                str(row["release_allowed"]),
            ]
        )
    add_table(
        document,
        ["Fit slot", "Risk band", "Manual tier", "Rows", "Release allowed"],
        locked_rows,
        widths=[2.2, 4.4, 4.0, 1.3, 2.0],
        font_size=6,
    )

    add_callout(
        document,
        "해석",
        "타자 쪽은 Savant 기반 성공/실패 classifier가 모두 pilot_promote로 올라왔기 때문에 후보명까지 제시할 근거가 있다. "
        "투수 쪽은 MiLB damage/command 모델이 아직 watch/diagnostic 단계라, 결과를 '후보 추천'보다 '추가 검증 우선순위'로 발표해야 한다. "
        "이 정직함이 오히려 모델 타당성을 높인다.",
        fill="EAF2FF",
    )
    document.add_page_break()


def add_final_candidates(document: Document, charts: dict[str, Path]) -> None:
    add_heading(document, "7. 최종 후보 및 후보별 통과 논리", 1)
    add_picture(document, charts["candidate_probabilities"], width_inches=6.5)

    top = read_csv("data_mining_recommendations_top3_v1.csv")
    hitters = read_csv("data_mining_hitter_candidates_v1.csv")
    pitchers = read_csv("data_mining_pitcher_candidates_v1.csv")

    hitter_rows = []
    for _, row in top[top["slot"].eq("foreign_hitter")].iterrows():
        detail = hitters[hitters["player_name"].eq(row["player_name"])].iloc[0]
        hitter_rows.append(
            [
                fmt_int(row["rank"]),
                row["player_name"],
                row["roster_team"],
                fmt_int(row["age"]),
                str(row["is_40man"]),
                fmt_pct(row["dm_success_prob"], 1),
                fmt_pct(row["dm_failure_prob"], 1),
                fmt_num(row["dm_margin"], 3),
                f"PA {fmt_int(detail['recent_pa'])}, wOBA {fmt_num(detail['recent_woba'], 3)}, K% {fmt_pct(detail['recent_k_pct'], 1)}, HH% {fmt_pct(detail['recent_hardhit_rate'], 1)}",
            ]
        )
    add_heading(document, "외인타자 Top 3", 2)
    add_table(
        document,
        ["Rank", "선수", "Org", "Age", "40-man", "P(success)", "P(failure)", "Margin", "최근 숫자"],
        hitter_rows,
        widths=[0.9, 2.7, 1.1, 1.0, 1.4, 1.6, 1.6, 1.3, 4.9],
        font_size=6,
    )

    add_callout(
        document,
        "타자 후보 해석",
        "Matos, Jones, Carlson은 단순 수동 점수표가 아니라 과거 KBO 외인타자 성공/실패 패턴을 학습한 ridge logistic classifier를 통과했다. "
        "공통점은 40인 제외 후보로 접근성 gate를 통과했고, 작은 샘플의 반짝 후보보다 안정적인 최근 PA와 Savant 입력 변수를 제공한다는 점이다.",
        fill="E8F5E9",
    )

    card_text = {
        "Luis Matos": [
            "모델상 1순위. P(success) 92.4%, P(failure) 8.2%, margin 0.842.",
            "최근 PA 205로 표본이 비교적 충분하고, 외야수/우타/24세라 upside와 포지션 fit이 동시에 존재한다.",
            "SSG fit 관점: 1루 주자 상황에서 필요한 contact floor와 batted-ball quality를 동시에 확인해야 한다.",
            "계약 gate: 저비용 가능성은 있으나 정확한 조직 통제/이적 의사 확인이 필수다.",
        ],
        "Nolan Jones": [
            "모델상 2순위. P(success) 90.2%, P(failure) 9.2%, margin 0.810.",
            "좌타 외야수로 SSG 외야/DH runway 문제와 연결 가능성이 있다.",
            "단, 공개 비용 신호가 약하지 않으므로 실제 취득 가능성이 가장 큰 반증 포인트다.",
            "계약 gate를 넘지 못하면 모델 점수와 별개로 탈락시켜야 한다.",
        ],
        "Dylan Carlson": [
            "모델상 3순위. P(success) 82.4%, P(failure) 15.1%, margin 0.673.",
            "양/좌우 대응성, 외야 포지션, 나이 면에서 후보 pool에 남길 이유가 있다.",
            "다만 현재 조직/계약 조건 소스가 흔들리므로 실무 검증 전에는 확정 후보로 과장하면 안 된다.",
            "타자 1, 2순위 대비 추천 강도는 낮지만, 모델이 뽑은 예비 3순위로 유지한다.",
        ],
    }
    for player, bullets in card_text.items():
        add_heading(document, player, 2)
        for item in bullets:
            add_bullet(document, item)

    pitcher_rows = []
    for _, row in top[top["slot"].eq("foreign_pitcher")].iterrows():
        detail = pitchers[pitchers["player_name"].eq(row["player_name"])].iloc[0]
        pitcher_rows.append(
            [
                fmt_int(row["rank"]),
                row["player_name"],
                row["roster_team"],
                fmt_int(row["age"]),
                str(row["is_40man"]),
                fmt_pct(row["dm_success_prob"], 1),
                fmt_pct(row["dm_failure_prob"], 1),
                fmt_num(row["dm_margin"], 3),
                f"MiLB IP {fmt_num(detail['pre_kbo_milb_ip'], 1)}, K/9 {fmt_num(detail['pre_kbo_milb_k9'], 2)}, BB/9 {fmt_num(detail['pre_kbo_milb_bb9'], 2)}, HR/9 {fmt_num(detail['pre_kbo_milb_hr9'], 2)}",
            ]
        )
    add_heading(document, "외인투수 진단 리드 Top 3", 2)
    add_table(
        document,
        ["Rank", "선수", "Org", "Age", "40-man", "P(success)", "P(failure)", "Margin", "최근 숫자"],
        pitcher_rows,
        widths=[0.9, 2.7, 1.1, 1.0, 1.4, 1.6, 1.6, 1.3, 5.1],
        font_size=6,
    )

    add_callout(
        document,
        "투수 후보 해석",
        "Wilson과 Gomber는 non-40man, MiLB IP 표본, damage/command feature gate를 통과한 진단 리드다. "
        "하지만 모델 검증력이 watch 수준이므로 '우리의 최종 추천 투수'라고 말하면 안 된다. "
        "Enns는 gate를 통과했지만 margin이 음수이고 공개 비용도 높아 hold로 처리하는 것이 맞다.",
        fill="FFF7D6",
    )

    pitcher_cards = {
        "Bryse Wilson": [
            "진단 리드 1순위. P(success) 50.5%, P(failure warning) 46.9%, margin 0.036.",
            "MiLB/최근 집계에서 103 IP, K/9 8.65, BB/9 2.79, HR/9 0.96으로 표본과 command/damage signal이 있다.",
            "SSG fit 관점: 선발 지속성, 80-90구, runner-on-base damage suppression을 영상/스플릿으로 확인해야 한다.",
            "현재 수준에서는 추천보다 추가 검증 우선순위다.",
        ],
        "Austin Gomber": [
            "진단 리드 2순위. P(success) 52.6%, P(failure warning) 49.2%, margin 0.034.",
            "172 IP, 37 GS에 가까운 workload continuity가 강점이다.",
            "SSG 메시지의 'load-bearing starter'와 가장 직관적으로 연결되지만, HR/9 1.46은 반드시 검증해야 한다.",
            "계약은 minor context로 보이나 정확한 비용/옵트아웃 확인 전 확정 금지다.",
        ],
        "Dietrich Enns": [
            "모델 gate는 통과했지만 margin -0.171로 hold다.",
            "K/9 10.34와 WHIP 1.28만 보면 끌릴 수 있으나, 모델은 실패 위험이 더 크다고 본다.",
            "공개 비용 신호도 높아 SSG 현실성에는 더 불리하다.",
            "따라서 Top 3 표에는 남기되 추천 후보가 아니라 반례/보류 후보로 설명한다.",
        ],
    }
    for player, bullets in pitcher_cards.items():
        add_heading(document, player, 2)
        for item in bullets:
            add_bullet(document, item)
    document.add_page_break()


def add_old_vs_new(document: Document) -> None:
    add_heading(document, "8. 이전 점수표와 데이터 마이닝 결론의 차이", 1)
    old = read_csv("unified_foreign_recommendations_top3_structured_only_v2.csv")
    old_rows = []
    for _, row in old.iterrows():
        old_rows.append(
            [
                row["slot_label"],
                fmt_int(row["recommendation_rank"]),
                row["player_name"],
                row["team_or_org"],
                fmt_num(row["unified_fit_score"], 2),
                row["data_mining_reason"],
            ]
        )
    add_table(
        document,
        ["Slot", "Rank", "기존 후보", "Org", "Score", "당시 이유"],
        old_rows,
        widths=[1.5, 1.0, 2.8, 1.2, 1.4, 8.0],
        font_size=6,
    )
    add_callout(
        document,
        "왜 후보가 바뀌었나",
        "기존 v2는 여러 feature block을 사람이 설계한 가중치로 합친 structured score였다. "
        "그 방식은 SSG fit을 잘 반영하지만, 표본이 작은 후보도 높은 점수를 받을 수 있다. "
        "새 data-mining v1은 과거 KBO 외인 성공/실패 classifier를 직접 통과시키므로 작은 표본 후보가 내려가고, "
        "Matos/Jones/Carlson처럼 모델이 낮은 실패확률로 보는 후보가 올라왔다.",
        fill="EAF2FF",
    )
    add_heading(document, "발표용 정리", 2)
    for item in [
        "우리는 처음에 SSG fit 기반 통합 점수표를 만들었다.",
        "그 다음 '이 점수가 실제 KBO 외인 성공/실패 패턴을 통과하는가?'를 검증했다.",
        "그 결과 타자 후보는 Matos, Jones, Carlson으로 재정렬되었다.",
        "투수는 아직 classifier가 약하므로 Fleming/Zimmermann/Yariel 같은 기존 score 후보와 Wilson/Gomber 같은 data-mining diagnostic 후보를 모두 재검증해야 한다.",
    ]:
        add_bullet(document, item)
    document.add_page_break()


def add_final_message_and_next_steps(document: Document) -> None:
    add_heading(document, "9. 최종 메시지와 다음 액션", 1)
    add_callout(
        document,
        "팀원 공유용 최종 문장",
        "SSG의 보강은 단순히 '장타 외야수 1명' 또는 '구속 좋은 투수 1명'이 아니다. "
        "타자는 1루 주자 상황을 득점권으로 번역하는 OF/DH runway bat, 투수는 초반 traffic과 볼넷/피홈런을 줄여 "
        "경기를 길게 끌고 가는 load-bearing starter가 핵심이다. 이 메시지를 과거 KBO 외인 성공/실패 데이터에 통과시킨 결과, "
        "타자 쪽은 Matos/Jones/Carlson을 우선 검토할 수 있고, 투수 쪽은 Wilson/Gomber를 진단 리드로 두되 추가 데이터 없이는 확정 추천하지 않는다.",
        fill="E8F5E9",
    )
    add_heading(document, "최종 후보 운영 방침", 2)
    for item in [
        "외인타자 1차 보드: Luis Matos, Nolan Jones, Dylan Carlson.",
        "외인투수 1차 진단 보드: Bryse Wilson, Austin Gomber; Dietrich Enns는 hold/negative-margin case.",
        "투수 최종 추천은 아직 보류: pitcher classifier는 watch 수준이라 추가 데이터와 수동 검토가 필수다.",
        "아시아쿼터는 이번 문서에서 후보명을 확정하지 않는다. 이유는 nationality/league-history/200k cap gate가 아직 후보 단위로 완전히 붙지 않았기 때문이다.",
    ]:
        add_bullet(document, item)

    add_heading(document, "남은 작업", 2)
    next_rows = [
        ["계약/연봉", "Top 20 후보 전체 current salary, guarantee, opt-out, buyout, assignment status 수집", "후보 제거 hard gate"],
        ["메디컬", "IL/rehab/workload drop/recent velocity decline를 구조화", "실패 리스크 모델 강화"],
        ["투수 모델", "과거 KBO 외인투수 입단 전 MiLB/MLB feature 추가 backfill", "watch -> pilot_promote 목표"],
        ["아시아쿼터", "NPB/CPBL 후보 nationality, prior/current Asian league eligibility, 200k cap 확인", "별도 shortlist 생성"],
        ["발표 자료", "Layer 1 메시지, 모델 검증, 후보 카드 3장으로 슬라이드화", "교수님/팀원 설명 가능"],
    ]
    add_table(
        document,
        ["영역", "해야 할 일", "목적"],
        next_rows,
        widths=[2.0, 8.7, 4.2],
        font_size=7,
    )
    document.add_page_break()


def add_appendix(document: Document) -> None:
    add_heading(document, "10. Appendix: 사용 산출물", 1)
    outputs = [
        "outputs/tables/data_coverage_by_league_v1.csv",
        "outputs/tables/project_data_coverage_audit_v1.csv",
        "outputs/tables/recruitment_gate_status_v33.csv",
        "outputs/tables/ssg_2026_runway_gap_by_team.csv",
        "outputs/tables/ssg_2026_role_runway_context.csv",
        "outputs/tables/ssg_hidden_weakness_robustness_decisions_v3.csv",
        "outputs/tables/ssg_hidden_weakness_final_message_v3.csv",
        "outputs/tables/ssg_pitching_message_v0_2_decision_table.csv",
        "outputs/tables/ssg_pitching_message_v0_2_context_validation.csv",
        "outputs/tables/ssg_2026_import_slot_pitching_impact.csv",
        "outputs/tables/kbo_foreign_archetype_prearrival_profile_v0_2.csv",
        "outputs/tables/kbo_foreign_failure_patterns_v0_1.csv",
        "outputs/tables/candidate_market_coverage_v0_3.csv",
        "outputs/tables/candidate_failure_risk_slot_summary_v0_1.csv",
        "outputs/tables/layer5_failure_risk_v0_3_slot_summary_v0_1.csv",
        "outputs/tables/kbo_translation_model_readiness_v0_3.csv",
        "outputs/tables/kbo_translation_failure_feature_family_decisions_v0_3.csv",
        "outputs/tables/data_mining_model_audit_v1.csv",
        "outputs/tables/data_mining_recommendations_top3_v1.csv",
        "outputs/tables/data_mining_hitter_candidates_v1.csv",
        "outputs/tables/data_mining_pitcher_candidates_v1.csv",
        "outputs/tables/kbo_contract_constraints_v1.csv",
        "outputs/tables/unified_foreign_recommendations_top3_structured_only_v2.csv",
        "scripts/build_data_mining_recommendation_model_v1.py",
    ]
    for output in outputs:
        add_bullet(document, output, size=8)

    add_heading(document, "공개 계약/연봉 확인 링크", 2)
    source_rows = [
        ["Spotrac Luis Matos / Brewers payroll", "https://www.spotrac.com/mlb/milwaukee-brewers/payroll/_/year/2026"],
        ["Spotrac Nolan Jones player contract", "https://www.spotrac.com/mlb/player/_/id/20565/nolan-jones"],
        ["Spotrac Cleveland cash/payroll", "https://www.spotrac.com/mlb/cleveland-guardians/cash"],
        ["Spotrac Phillies payroll / Bryse Wilson", "https://www.spotrac.com/mlb/philadelphia-phillies/payroll/_/year/2026"],
        ["MiLB Austin Gomber player page", "https://www.milb.com/player/austin-gomber-596295"],
        ["Spotrac Austin Gomber contract", "https://www.spotrac.com/mlb/player/_/id/24562/austin-gomber"],
        ["Spotrac Orioles overview / Dietrich Enns", "https://www.spotrac.com/mlb/baltimore-orioles/overview"],
        ["Dylan Carlson salary signal", "https://www.spotrac.com/mlb/rankings/player/_/year/2026/team/chc/position/of/sort/cap_total"],
    ]
    add_table(document, ["Source", "URL"], source_rows, widths=[5.0, 10.0], font_size=6)

    add_callout(
        document,
        "문서 한계",
        "이 보고서는 현재 확보된 숫자형/구조화 데이터 기준의 정식 분석본이다. 실제 영입 가능성은 구단 간 권리, 선수 의향, "
        "보장계약, opt-out, buyout, 메디컬, 비자/가족 변수에 따라 바뀐다. 따라서 이 문서는 final signing answer가 아니라 "
        "데이터 마이닝 기반 shortlist 및 검증 우선순위 보고서로 사용하는 것이 맞다.",
        fill="F3F4F6",
    )


def build_doc() -> Path:
    DOC_OUT_DIR.mkdir(parents=True, exist_ok=True)
    charts = make_charts()
    document = Document()
    apply_doc_styles(document)

    add_title_page(document)
    add_table_of_contents(document)
    add_data_and_model_design(document, charts)
    add_ssg_hitting_layer(document, charts)
    add_ssg_pitching_layer(document, charts)
    add_kbo_archetype_layer(document)
    add_market_and_contract_layer(document)
    add_translation_and_risk_layer(document)
    add_final_candidates(document, charts)
    add_old_vs_new(document)
    add_final_message_and_next_steps(document)
    add_appendix(document)

    section = document.add_section(WD_SECTION.CONTINUOUS)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    document.save(DOC_PATH)
    return DOC_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
